import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
import google.generativeai as genai
from docx import Document
from PIL import Image
from dotenv import load_dotenv
import threading

# --- LOAD KONFIGURASI ---
load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")

if api_key:
    genai.configure(api_key=api_key)

class GeminiOCRApp:
    def __init__(self, root):
        self.root = root
        self.root.title("AI Document Digitizer Pro")
        self.root.geometry("550x450")
        self.root.configure(bg="#ffffff")
        
        self.style = ttk.Style()
        self.style.theme_use('clam')
        self.style.configure("TButton", font=("Segoe UI", 10), padding=6)
        self.style.configure("Action.TButton", background="#2b579a", foreground="white")

        self.file_path = ""

        # --- UI LAYOUT ---
        self.main_container = tk.Frame(root, bg="#ffffff", padx=30, pady=30)
        self.main_container.pack(expand=True, fill="both")

        tk.Label(self.main_container, text="Gemini OCR Tool", font=("Segoe UI", 18, "bold"), bg="#ffffff", fg="#333333").pack(pady=(0, 5))
        tk.Label(self.main_container, text="Ekstrak Gambar ke Tabel Word Asli", font=("Segoe UI", 9), bg="#ffffff", fg="#777777").pack(pady=(0, 25))

        self.file_frame = tk.Frame(self.main_container, bg="#f8f9fa", bd=1, relief="solid")
        self.file_frame.pack(fill="x", pady=10)

        self.lbl_filename = tk.Label(self.file_frame, text="Belum ada file dipilih", font=("Segoe UI", 9, "italic"), bg="#f8f9fa", fg="#999999", padx=10, pady=15)
        self.lbl_filename.pack(side="left", expand=True)

        self.btn_select = ttk.Button(self.main_container, text="Pilih Gambar", command=self.select_file)
        self.btn_select.pack(pady=10)

        tk.Frame(self.main_container, height=1, bg="#eeeeee").pack(fill="x", pady=20)

        self.btn_process = ttk.Button(self.main_container, text="Mulai Ekstraksi ke Word", style="Action.TButton", command=self.start_processing_thread)
        self.btn_process.pack(fill="x", pady=5)

        self.status_var = tk.StringVar(value="Ready")
        self.lbl_status = tk.Label(self.main_container, textvariable=self.status_var, font=("Segoe UI", 8), bg="#ffffff", fg="#2b579a")
        self.lbl_status.pack(pady=5)

    def select_file(self):
        file_selected = filedialog.askopenfilename(filetypes=[("Images", "*.jpg *.jpeg *.png *.webp")])
        if file_selected:
            self.file_path = file_selected
            self.lbl_filename.config(text=os.path.basename(self.file_path), fg="#333333", font=("Segoe UI", 9, "bold"))
            self.status_var.set("File siap diproses")

    def start_processing_thread(self):
        if not self.file_path:
            messagebox.showwarning("Peringatan", "Silakan pilih gambar terlebih dahulu.")
            return
        self.btn_process.state(['disabled'])
        self.btn_select.state(['disabled'])
        self.status_var.set("Sedang memproses (mohon tunggu)...")
        threading.Thread(target=self.process_ocr, daemon=True).start()

    def add_content_to_docx(self, doc, text_result):
        """Mendeteksi tabel Markdown dan mengonversinya ke tabel Word asli"""
        lines = text_result.strip().split('\n')
        i = 0
        
        # Daftar kata kunci yang sering muncul di screenshot UI Word (untuk diabaikan)
        noise_keywords = ["AutoSave", "Protected View", "Enable Editing", "File", "Home", "Insert", "Layout", "References"]

        while i < len(lines):
            line = lines[i].strip()
            
            # Deteksi awal tabel Markdown (baris yang mengandung setidaknya dua karakter '|')
            if line.count('|') >= 2:
                table_data = []
                while i < len(lines) and lines[i].strip().count('|') >= 2:
                    raw_line = lines[i].strip()
                    # Lewati baris separator seperti |---|---|
                    if not all(c in '| -:' for c in raw_line):
                        # Pecah sel berdasarkan '|' dan bersihkan spasi
                        cells = [c.strip() for c in raw_line.split('|') if c.strip()]
                        if cells:
                            table_data.append(cells)
                    i += 1
                
                if table_data:
                    # Buat tabel asli di Word
                    rows = len(table_data)
                    cols = max(len(row) for row in table_data)
                    table = doc.add_table(rows=rows, cols=cols)
                    table.style = 'Table Grid' # Memberikan border
                    
                    for r_idx, row_data in enumerate(table_data):
                        for c_idx, cell_value in enumerate(row_data):
                            if c_idx < cols:
                                table.cell(r_idx, c_idx).text = cell_value
                    doc.add_paragraph() # Jarak setelah tabel
                continue # Lanjut ke loop berikutnya tanpa menambah i lagi
            
            # Jika teks biasa (bukan tabel) dan bukan merupakan menu/UI Word
            if line and not any(noise in line for noise in noise_keywords):
                doc.add_paragraph(line)
            
            i += 1

    def process_ocr(self):
        try:
            model = genai.GenerativeModel('gemini-2.5-flash')
            img = Image.open(self.file_path)

            prompt = (
                "Ekstrak teks dari gambar ini dengan sangat teliti. "
                "PENTING: Jika ada tabel, buatlah dalam format tabel Markdown. "
                "Jangan sertakan elemen UI software seperti nama menu atau status bar. "
                "Langsung berikan hasil ekstraksi saja tanpa basa-basi."
            )
            
            response = model.generate_content([prompt, img])
            text_result = response.text

            doc = Document()
            doc.add_heading('Hasil Konversi Formulir', 0)
            
            self.add_content_to_docx(doc, text_result)
            
            self.root.after(0, self.save_document, doc)

        except Exception as err:
            self.root.after(0, lambda e=err: messagebox.showerror("Error", f"Gagal memproses: {str(e)}"))
        finally:
            self.root.after(0, self.reset_ui)

    def save_document(self, doc):
        save_path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=os.path.splitext(os.path.basename(self.file_path))[0] + "_hasil"
        )
        if save_path:
            doc.save(save_path)
            messagebox.showinfo("Sukses", "Dokumen berhasil disimpan dengan tabel asli!")

    def reset_ui(self):
        self.btn_process.state(['!disabled'])
        self.btn_select.state(['!disabled'])
        self.status_var.set("Ready")

if __name__ == "__main__":
    root = tk.Tk()
    app = GeminiOCRApp(root)
    root.mainloop()