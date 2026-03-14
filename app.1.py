import os
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from google import genai
from PIL import Image, ImageTk
import threading
import json
import re
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.section import WD_ORIENT

CONFIG_FILE = "config_docx.json"

PROMPT_MASTER = """
Kamu adalah ahli rekonstruksi dokumen. Analisis gambar dokumen ini dengan sangat teliti dan rekonstruksi ke format JSON yang presisi.

TUGAS UTAMA:
Ekstrak SEMUA elemen teks, tabel, dan struktur layout dengan akurasi tinggi.

FORMAT JSON WAJIB (kembalikan HANYA JSON, tanpa teks lain):
{
  "page_layout": {
    "orientation": "portrait" atau "landscape",
    "has_letterhead": true/false
  },
  "elements": [
    {
      "type": "paragraph" | "heading" | "list_item" | "table" | "field_label" | "field_value" | "signature_block",
      "content": "teks isi",
      "alignment": "left" | "center" | "right" | "justify",
      "bold": true/false,
      "italic": true/false,
      "underline": true/false,
      "font_size": 8-24 (estimasi dalam pt, default 11),
      "level": 1-6 (hanya untuk heading),
      "indent_level": 0-4 (indentasi paragraf),
      "space_before": 0-24 (spasi sebelum dalam pt),
      "space_after": 0-24 (spasi sesudah dalam pt),
      "is_uppercase": true/false,
      "list_type": "bullet" | "number" | null,
      "items": ["item1", "item2"] (hanya untuk list),
      "rows": [["col1","col2"],["val1","val2"]] (hanya untuk table),
      "has_header_row": true/false (untuk table),
      "col_widths_pct": [50, 50] (persentase lebar kolom untuk table),
      "border": true/false (untuk table)
    }
  ]
}

ATURAN EKSTRAKSI:
1. Baca SETIAP karakter teks dengan teliti - jangan lewatkan satu pun
2. Untuk FORMULIR: identifikasi label fields (misal "Nama:", "Tanggal:") dan nilainya sebagai field_label / field_value
3. Untuk TABEL: ekstrak semua baris dan kolom secara lengkap, perkirakan lebar kolom relatif
4. Untuk HEADING/JUDUL: identifikasi ukuran font relatif (judul besar = 16-18pt, subjudul = 13-14pt, normal = 11pt)
5. Teks UPPERCASE biasanya adalah judul formal - set is_uppercase: true
6. Perhatikan spasi antar paragraf dan indentasi
7. Teks di tengah halaman (seperti kop surat, judul) harus alignment: "center"
8. Teks paragraf panjang biasanya alignment: "justify"
9. Abaikan coretan, tulisan yang tidak terbaca, atau artefak scan
10. Untuk blok tanda tangan: gunakan type "signature_block"
11. Perhatikan apakah ada teks tebal (bold), miring (italic), atau garis bawah (underline)

PENTING SEKALI:
- Kembalikan HANYA JSON yang valid. Tidak ada penjelasan, tidak ada markdown code block.
- Semua nilai string dalam JSON HARUS menggunakan tanda kutip ganda (").
- Karakter kutip ganda di dalam string harus di-escape dengan backslash: \\"
- Karakter backslash di dalam string harus di-escape: \\\\
- Jangan gunakan newline di dalam nilai string JSON; ganti dengan spasi.
"""

AVAILABLE_MODELS = [
    "gemini-2.5-flash",
    "gemini-2.5-pro",
    "gemini-2.0-flash",
    "gemini-1.5-flash",
    "gemini-1.5-pro",
]

PAGE_CONTENT_WIDTH_TWIP_PORTRAIT  = 9026
PAGE_CONTENT_WIDTH_TWIP_LANDSCAPE = 14798


def sanitize_filename(name: str) -> str:
    name = re.sub(r'[\\/:*?"<>|]', '_', name)
    name = re.sub(r'\s+', '_', name.strip())
    name = re.sub(r'_+', '_', name)
    return name.strip('_')


def auto_increment_path(path: str) -> str:
    if not os.path.exists(path):
        return path
    base, ext = os.path.splitext(path)
    base = re.sub(r'_v\d+$', '', base)
    version = 2
    while True:
        candidate = f"{base}_v{version}{ext}"
        if not os.path.exists(candidate):
            return candidate
        version += 1


# ══════════════════════════════════════════════════════════════════
# ROBUST JSON PARSER — KUNCI PERBAIKAN ERROR
# ══════════════════════════════════════════════════════════════════

def extract_and_fix_json(raw_text: str) -> dict:
    """
    Mencoba berbagai strategi untuk mengekstrak JSON valid dari respons AI.
    
    Strategi (berurutan):
    1. Parse langsung setelah bersihkan markdown fences
    2. Ekstrak blok {...} terluar
    3. Perbaiki masalah umum (trailing commas, kutip tunggal, newline dalam string)
    4. Gunakan json_repair jika tersedia
    5. Fallback: kembalikan struktur minimal
    """
    if not raw_text or not raw_text.strip():
        raise ValueError("Respons AI kosong")

    # ── Langkah 1: Bersihkan markdown fences ──
    cleaned = raw_text.strip()
    # Hapus ```json ... ``` atau ``` ... ```
    cleaned = re.sub(r'^```(?:json)?\s*', '', cleaned, flags=re.MULTILINE)
    cleaned = re.sub(r'\s*```\s*$', '', cleaned, flags=re.MULTILINE)
    cleaned = cleaned.strip().strip('`').strip()

    # ── Langkah 2: Coba parse langsung ──
    try:
        return json.loads(cleaned)
    except json.JSONDecodeError:
        pass

    # ── Langkah 3: Ekstrak blok JSON terluar {...} ──
    json_match = re.search(r'\{[\s\S]*\}', cleaned)
    if json_match:
        candidate = json_match.group(0)
        try:
            return json.loads(candidate)
        except json.JSONDecodeError:
            cleaned = candidate  # lanjutkan dengan kandidat ini

    # ── Langkah 4: Perbaiki masalah umum ──
    fixed = _repair_json_string(cleaned)
    try:
        return json.loads(fixed)
    except json.JSONDecodeError:
        pass

    # ── Langkah 5: Coba json_repair (library opsional) ──
    try:
        from json_repair import repair_json
        repaired = repair_json(cleaned)
        result = json.loads(repaired)
        return result
    except (ImportError, Exception):
        pass

    # ── Langkah 6: Fallback — kembalikan struktur minimal ──
    # Simpan raw untuk debugging
    _save_debug_response(raw_text)
    raise json.JSONDecodeError(
        f"Tidak dapat memperbaiki JSON setelah semua strategi dicoba. "
        f"Raw response disimpan ke 'debug_last_response.txt'.\n"
        f"Potongan awal: {raw_text[:300]}",
        cleaned, 0
    )


def _repair_json_string(text: str) -> str:
    """Perbaikan heuristik untuk JSON yang hampir valid."""

    # 1) Hapus trailing commas sebelum } atau ]
    #    Contoh: {"a": 1,} atau ["x",]
    text = re.sub(r',\s*([}\]])', r'\1', text)

    # 2) Ganti newline di dalam string JSON value dengan spasi
    #    (Newline literal di dalam string tidak valid di JSON)
    #    Strategi: temukan semua string "..." dan bersihkan isinya
    def clean_string_value(m):
        inner = m.group(1)
        # Ganti newline literal dengan \n (escaped)
        inner = inner.replace('\n', ' ').replace('\r', ' ')
        # Ganti tab literal dengan spasi
        inner = inner.replace('\t', ' ')
        return f'"{inner}"'

    # Regex sederhana untuk string JSON (tidak nested)
    text = re.sub(r'"((?:[^"\\]|\\.)*)"', clean_string_value, text)

    # 3) Perbaiki boolean/null yang salah kapital
    text = re.sub(r'\bTrue\b',  'true',  text)
    text = re.sub(r'\bFalse\b', 'false', text)
    text = re.sub(r'\bNone\b',  'null',  text)
    text = re.sub(r'\bNull\b',  'null',  text)

    # 4) Hapus komentar // ... (tidak valid di JSON)
    text = re.sub(r'//[^\n]*', '', text)

    # 5) Perbaiki single-quote yang digunakan sebagai pembatas string
    #    Hanya jika tidak ada double-quote sama sekali di area itu
    #    (Hati-hati: apostrof dalam teks Indonesia jangan diubah)
    # Strategi konservatif: skip jika ada banyak double-quote
    if text.count('"') < text.count("'") * 0.5:
        text = re.sub(r"'([^']*)'", r'"\1"', text)

    return text


def _save_debug_response(raw: str):
    """Simpan respons mentah AI untuk debugging."""
    try:
        debug_path = "debug_last_response.txt"
        with open(debug_path, "w", encoding="utf-8") as f:
            f.write(raw)
        print(f"[DEBUG] Raw AI response disimpan ke: {debug_path}")
    except Exception:
        pass


# ══════════════════════════════════════════════════════════════════
# HELPER TABEL
# ══════════════════════════════════════════════════════════════════

def _set_cell_shading(cell, fill_hex: str):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    for old in tcPr.findall(qn('w:shd')):
        tcPr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  fill_hex.upper())
    tcPr.append(shd)


def _set_cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in (('top', top), ('bottom', bottom), ('left', left), ('right', right)):
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    for old in tcPr.findall(qn('w:tcMar')):
        tcPr.remove(old)
    tcPr.append(tcMar)


def _set_cell_width(cell, width_twip: int):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcW  = OxmlElement('w:tcW')
    tcW.set(qn('w:w'),    str(int(width_twip)))
    tcW.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')):
        tcPr.remove(old)
    tcPr.append(tcW)


def _set_table_width(table, width_twip: int):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblW = OxmlElement('w:tblW')
    tblW.set(qn('w:w'),    str(int(width_twip)))
    tblW.set(qn('w:type'), 'dxa')
    for old in tblPr.findall(qn('w:tblW')):
        tblPr.remove(old)
    tblPr.append(tblW)


def _set_table_layout_fixed(table):
    tbl  = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    for old in tblPr.findall(qn('w:tblLayout')):
        tblPr.remove(old)
    tblPr.append(tblLayout)


def _set_col_widths(table, col_widths_twip: list):
    tbl = table._tbl
    tblGrid = tbl.find(qn('w:tblGrid'))
    if tblGrid is None:
        tblGrid = OxmlElement('w:tblGrid')
        tbl.insert(1, tblGrid)
    else:
        for gc in list(tblGrid):
            tblGrid.remove(gc)
    for w in col_widths_twip:
        gridCol = OxmlElement('w:gridCol')
        gridCol.set(qn('w:w'), str(int(w)))
        tblGrid.append(gridCol)
    for row in table.rows:
        cells = row.cells
        for i, cell in enumerate(cells):
            if i < len(col_widths_twip):
                _set_cell_width(cell, col_widths_twip[i])


def build_table(doc, el: dict, content_width_twip: int) -> None:
    rows_data = el.get('rows', [])
    if not rows_data:
        return
    num_cols = max(len(r) for r in rows_data)
    if num_cols == 0:
        return
    rows_data = [list(r) + [''] * (num_cols - len(r)) for r in rows_data]

    col_pct = el.get('col_widths_pct', [])
    if col_pct and len(col_pct) == num_cols:
        total_pct = sum(col_pct) or 100
        col_widths_twip = [int(content_width_twip * p / total_pct) for p in col_pct]
    else:
        unit = content_width_twip // num_cols
        col_widths_twip = [unit] * num_cols

    diff = content_width_twip - sum(col_widths_twip)
    col_widths_twip[-1] += diff

    has_border = el.get('border', True)
    try:
        table = doc.add_table(rows=len(rows_data), cols=num_cols)
        table.style = 'Table Grid' if has_border else 'Normal Table'
    except KeyError:
        table = doc.add_table(rows=len(rows_data), cols=num_cols)

    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_table_width(table, content_width_twip)
    _set_col_widths(table, col_widths_twip)
    _set_table_layout_fixed(table)

    has_header = el.get('has_header_row', False)
    font_size  = el.get('font_size', 10)
    align_map  = {
        "left":    WD_ALIGN_PARAGRAPH.LEFT,
        "center":  WD_ALIGN_PARAGRAPH.CENTER,
        "right":   WD_ALIGN_PARAGRAPH.RIGHT,
        "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    cell_align = align_map.get(el.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)

    for r_idx, row_data in enumerate(rows_data):
        row_obj = table.rows[r_idx]
        is_header_row = has_header and r_idx == 0
        for c_idx in range(num_cols):
            cell_text = str(row_data[c_idx]) if row_data[c_idx] is not None else ""
            cell = row_obj.cells[c_idx]
            _set_cell_margins(cell, top=80, bottom=80, left=120, right=120)
            _set_cell_width(cell, col_widths_twip[c_idx])
            if is_header_row:
                _set_cell_shading(cell, "D9E1F2")
            para = cell.paragraphs[0]
            para.alignment = cell_align
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after  = Pt(0)
            run = para.add_run(cell_text)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(font_size)
            if is_header_row:
                run.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(6)


# ══════════════════════════════════════════════════════════════════
# APLIKASI UTAMA
# ══════════════════════════════════════════════════════════════════

class OmbudsmanIntelligentDocx:
    def __init__(self, root):
        self.root = root
        self.root.title("Ombudsman Master Reconstructor v3.4")
        self.root.geometry("600x580")
        self.root.configure(bg="#f5f6fa")

        self.api_key    = None
        self.model_name = "gemini-2.5-flash"
        self.client     = None
        self.file_paths = []

        self.check_setup()
        self.setup_ui()

    # ──────────────────────────────────────────────────────────────
    # CONFIG / SETUP
    # ──────────────────────────────────────────────────────────────

    def check_setup(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE, "r") as f:
                    config = json.load(f)
                    self.api_key    = config.get("api_key")
                    self.model_name = config.get("model_name", "gemini-2.5-flash")
            except Exception:
                pass
        if not self.api_key:
            self.run_first_time_setup()
        else:
            self.init_genai_client()

    def save_config(self):
        config = {"api_key": self.api_key, "model_name": self.model_name}
        with open(CONFIG_FILE, "w") as f:
            json.dump(config, f, indent=2)

    def init_genai_client(self):
        try:
            self.client = genai.Client(api_key=self.api_key)
        except Exception as e:
            messagebox.showerror("Error", f"Gagal inisialisasi API: {str(e)}")

    def run_first_time_setup(self):
        setup_win = tk.Toplevel(self.root)
        setup_win.title("Setup API — Pertama Kali")
        setup_win.geometry("440x220")
        setup_win.grab_set()
        main_frame = tk.Frame(setup_win, padx=20, pady=20)
        main_frame.pack(expand=True, fill="both")

        tk.Label(main_frame, text="Gemini API Key:", font=("Segoe UI", 10)).pack(anchor="w")
        entry_key = ttk.Entry(main_frame, width=52, show="*")
        entry_key.pack(pady=5, fill="x")

        tk.Label(main_frame, text="Model:", font=("Segoe UI", 10)).pack(anchor="w", pady=(8, 0))
        model_var   = tk.StringVar(value=self.model_name)
        model_combo = ttk.Combobox(main_frame, textvariable=model_var,
                                   values=AVAILABLE_MODELS, state="readonly")
        model_combo.pack(fill="x")

        def save():
            key = entry_key.get().strip()
            if not key:
                messagebox.showwarning("Peringatan", "API key tidak boleh kosong!", parent=setup_win)
                return
            self.api_key    = key
            self.model_name = model_var.get()
            self.save_config()
            self.init_genai_client()
            setup_win.destroy()

        ttk.Button(main_frame, text="Simpan & Mulai", command=save).pack(pady=15)

    # ──────────────────────────────────────────────────────────────
    # SETTINGS PANEL
    # ──────────────────────────────────────────────────────────────

    def open_settings(self):
        win = tk.Toplevel(self.root)
        win.title("⚙ Pengaturan API")
        win.geometry("460x280")
        win.resizable(False, False)
        win.grab_set()
        win.configure(bg="#f5f6fa")

        frame = tk.Frame(win, bg="#ffffff", padx=24, pady=20)
        frame.pack(expand=True, fill="both", padx=16, pady=16)

        tk.Label(frame, text="Pengaturan API & Model",
                 font=("Segoe UI", 12, "bold"), bg="#ffffff", fg="#1a365d").grid(
            row=0, column=0, columnspan=2, sticky="w", pady=(0, 16))

        tk.Label(frame, text="Gemini API Key:", font=("Segoe UI", 10),
                 bg="#ffffff").grid(row=1, column=0, sticky="w", pady=4)
        entry_key = ttk.Entry(frame, width=36, show="*")
        entry_key.insert(0, self.api_key or "")
        entry_key.grid(row=1, column=1, sticky="ew", padx=(8, 0), pady=4)

        show_var = tk.BooleanVar(value=False)
        def toggle_show():
            entry_key.config(show="" if show_var.get() else "*")
        ttk.Checkbutton(frame, text="Tampilkan key", variable=show_var,
                        command=toggle_show).grid(row=2, column=1, sticky="w", padx=(8, 0))

        tk.Label(frame, text="Model AI:", font=("Segoe UI", 10),
                 bg="#ffffff").grid(row=3, column=0, sticky="w", pady=(12, 4))
        model_var   = tk.StringVar(value=self.model_name)
        model_combo = ttk.Combobox(frame, textvariable=model_var,
                                   values=AVAILABLE_MODELS, state="readonly", width=34)
        model_combo.grid(row=3, column=1, sticky="ew", padx=(8, 0), pady=(12, 4))

        model_hint = {
            "gemini-2.5-flash": "⚡ Cepat & hemat kuota — cocok untuk sehari-hari",
            "gemini-2.5-pro":   "🏆 Akurasi tertinggi — disarankan untuk dokumen kompleks",
            "gemini-2.0-flash": "⚡ Generasi terbaru, cepat",
            "gemini-1.5-flash": "⚡ Cepat, kuota lebih besar",
            "gemini-1.5-pro":   "🔎 Akurasi tinggi generasi sebelumnya",
        }
        hint_var = tk.StringVar(value=model_hint.get(self.model_name, ""))
        hint_label = tk.Label(frame, textvariable=hint_var, font=("Segoe UI", 8),
                              bg="#ffffff", fg="#718096", wraplength=280, justify="left")
        hint_label.grid(row=4, column=1, sticky="w", padx=(8, 0))

        def on_model_change(event=None):
            hint_var.set(model_hint.get(model_var.get(), ""))
        model_combo.bind("<<ComboboxSelected>>", on_model_change)
        frame.columnconfigure(1, weight=1)

        btn_frame = tk.Frame(win, bg="#f5f6fa")
        btn_frame.pack(fill="x", padx=16, pady=(0, 16))

        def apply_settings():
            key = entry_key.get().strip()
            if not key:
                messagebox.showwarning("Peringatan", "API key tidak boleh kosong!", parent=win)
                return
            changed_key = (key != self.api_key)
            self.api_key    = key
            self.model_name = model_var.get()
            self.save_config()
            if changed_key:
                self.init_genai_client()
            self.model_indicator_var.set(f"Model: {self.model_name}")
            self.update_status(f"Pengaturan disimpan. Model: {self.model_name}")
            win.destroy()

        ttk.Button(btn_frame, text="💾 Simpan", command=apply_settings).pack(side="right", padx=(8, 0))
        ttk.Button(btn_frame, text="Batal",     command=win.destroy).pack(side="right")

    # ──────────────────────────────────────────────────────────────
    # UI
    # ──────────────────────────────────────────────────────────────

    def setup_ui(self):
        style = ttk.Style()
        style.theme_use('clam')
        style.configure("Accent.TButton",   font=("Segoe UI", 10, "bold"),
                        foreground="white", background="#1a365d", padding=10)
        style.map("Accent.TButton", background=[('active', '#2c5282')])
        style.configure("Sub.TButton",      font=("Segoe UI", 9),
                        foreground="#1a365d", background="#e2e8f0", padding=8)
        style.configure("Settings.TButton", font=("Segoe UI", 9),
                        foreground="#4a5568", background="#edf2f7", padding=6)

        self.container = tk.Frame(self.root, bg="#ffffff", padx=30, pady=20)
        self.container.pack(expand=True, fill="both", padx=20, pady=20)
        self.container.columnconfigure(0, weight=1)

        # Header
        header_frame = tk.Frame(self.container, bg="#ffffff")
        header_frame.grid(row=0, column=0, sticky="ew", pady=(0, 16))
        header_frame.columnconfigure(1, weight=1)

        try:
            img_left = Image.open("logo-kiri1.jpg").resize((120, 120), Image.LANCZOS)
            self.logo_left = ImageTk.PhotoImage(img_left)
            tk.Label(header_frame, image=self.logo_left, bg="#ffffff").grid(row=0, column=0, sticky="w")
            img_right = Image.open("logo-kanan.png").resize((80, 80), Image.LANCZOS)
            self.logo_right = ImageTk.PhotoImage(img_right)
            tk.Label(header_frame, image=self.logo_right, bg="#ffffff").grid(row=0, column=2, sticky="e")
        except Exception as e:
            print(f"Informasi: Logo tidak dimuat ({e})")

        tk.Label(self.container, text="Document Reconstructor v3.4",
                 font=("Segoe UI", 16, "bold"), bg="#ffffff", fg="#1a365d").grid(row=1, column=0)

        self.status_var = tk.StringVar(value="● Status: Sistem Siap")
        self.status_label = tk.Label(self.container, textvariable=self.status_var,
                                     font=("Segoe UI", 9), bg="#ffffff", fg="#4a5568")
        self.status_label.grid(row=2, column=0, pady=(0, 2))

        self.model_indicator_var = tk.StringVar(value=f"Model: {self.model_name}")
        tk.Label(self.container, textvariable=self.model_indicator_var,
                 font=("Segoe UI", 8, "italic"), bg="#ffffff", fg="#a0aec0").grid(row=3, column=0, pady=(0, 8))

        self.progress = ttk.Progressbar(self.container, mode='indeterminate', length=400)
        self.progress.grid(row=4, column=0, sticky="ew", pady=(0, 12))

        list_frame = tk.Frame(self.container, bg="#ffffff")
        list_frame.grid(row=5, column=0, sticky="ew", pady=(0, 8))
        list_frame.columnconfigure(0, weight=1)
        self.file_listbox = tk.Listbox(list_frame, height=5, font=("Segoe UI", 8),
                                        bg="#f7fafc", relief="flat", borderwidth=1,
                                        selectmode=tk.EXTENDED)
        self.file_listbox.grid(row=0, column=0, sticky="ew")
        scrollbar = ttk.Scrollbar(list_frame, orient="vertical", command=self.file_listbox.yview)
        scrollbar.grid(row=0, column=1, sticky="ns")
        self.file_listbox.config(yscrollcommand=scrollbar.set)

        info_frame = tk.Frame(self.container, bg="#e8f4fd", bd=0)
        info_frame.grid(row=6, column=0, sticky="ew", pady=(0, 10))
        info_frame.columnconfigure(0, weight=1)

        prev_frame = tk.Frame(info_frame, bg="#dbeafe", padx=10, pady=4)
        prev_frame.grid(row=2, column=0, sticky="ew")
        prev_frame.columnconfigure(1, weight=1)
        tk.Label(prev_frame, text="📄 Preview:", font=("Segoe UI", 8, "bold"),
                 bg="#dbeafe", fg="#1e3a8a").grid(row=0, column=0, sticky="w", padx=(0, 6))
        self.preview_var = tk.StringVar(value="— Pilih file terlebih dahulu —")
        tk.Label(prev_frame, textvariable=self.preview_var,
                 font=("Segoe UI", 8), bg="#dbeafe", fg="#1e3a8a",
                 anchor="w").grid(row=0, column=1, sticky="ew")

        button_frame = tk.Frame(self.container, bg="#ffffff")
        button_frame.grid(row=7, column=0, sticky="ew")
        button_frame.columnconfigure(0, weight=1)
        button_frame.columnconfigure(1, weight=1)

        ttk.Button(button_frame, text="📁 PILIH FILE",
                   command=self.select_files, style="Accent.TButton").grid(
            row=0, column=0, sticky="ew", padx=(0, 5), pady=5)
        ttk.Button(button_frame, text="🗑 HAPUS PILIHAN",
                   command=self.clear_files, style="Sub.TButton").grid(
            row=0, column=1, sticky="ew", padx=(5, 0), pady=5)

        self.btn_process = ttk.Button(button_frame, text="🚀 KONVERT KE DOCX",
                                      command=self.start_processing, style="Accent.TButton")
        self.btn_process.grid(row=1, column=0, columnspan=2, sticky="ew", pady=5)

        ttk.Button(button_frame, text="⚙ Pengaturan API",
                   command=self.open_settings, style="Settings.TButton").grid(
            row=2, column=0, columnspan=2, sticky="ew", pady=(2, 5))

        self.container.rowconfigure(8, weight=1)
        footer_frame = tk.Frame(self.container, bg="#ffffff")
        footer_frame.grid(row=9, column=0, sticky="s", pady=(16, 0))
        tk.Frame(footer_frame, height=1, width=400, bg="#edf2f7").pack(pady=(0, 8))
        footer_text = ("© 2026 Ombudsman Republik Indonesia | Universitas Halu Oleo\n"
                       "Abrar Wujedan & Abdul Mu'iz Azizul Raeba")
        tk.Label(footer_frame, text=footer_text, font=("Segoe UI", 7, "bold"),
                 bg="#ffffff", fg="#94a3b8", justify="center").pack()

    # ──────────────────────────────────────────────────────────────
    # FILE NAMING
    # ──────────────────────────────────────────────────────────────

    def _build_auto_filename(self) -> str:
        if self.file_paths:
            base  = os.path.splitext(os.path.basename(self.file_paths[0]))[0]
            clean = sanitize_filename(base)
            if len(self.file_paths) > 1:
                clean = f"{clean}_dan_{len(self.file_paths)-1}_lainnya"
        else:
            clean = "Dokumen"
        date_str = datetime.now().strftime("%Y%m%d")
        return f"Rekonstruksi_{clean}_{date_str}"

    def _get_output_dir(self) -> str:
        if self.file_paths:
            return os.path.dirname(os.path.abspath(self.file_paths[0]))
        return os.path.expanduser("~\\Documents" if os.name == "nt" else "~/Documents")

    def _refresh_preview(self):
        name      = self._build_auto_filename()
        out_dir   = self._get_output_dir()
        candidate = auto_increment_path(os.path.join(out_dir, name + ".docx"))
        display   = os.path.basename(candidate)
        self.preview_var.set(f"{display}  (di {out_dir})")

    def _get_final_save_path(self) -> str:
        name    = self._build_auto_filename()
        out_dir = self._get_output_dir()
        os.makedirs(out_dir, exist_ok=True)
        path = os.path.join(out_dir, name + ".docx")
        return auto_increment_path(path)

    # ──────────────────────────────────────────────────────────────
    # FILE HELPERS
    # ──────────────────────────────────────────────────────────────

    def select_files(self):
        files = filedialog.askopenfilenames(
            filetypes=[("Gambar & PDF", "*.jpg *.jpeg *.png *.pdf *.tiff *.bmp")]
        )
        if files:
            for f in files:
                if f not in self.file_paths:
                    self.file_paths.append(f)
                    self.file_listbox.insert(tk.END, os.path.basename(f))
            self._refresh_preview()
            self.status_var.set(f"● {len(self.file_paths)} file dipilih.")

    def clear_files(self):
        self.file_paths.clear()
        self.file_listbox.delete(0, tk.END)
        self.preview_var.set("— Pilih file terlebih dahulu —")
        self.status_var.set("● Status: Sistem Siap")

    # ──────────────────────────────────────────────────────────────
    # PROCESSING — DENGAN RETRY & ROBUST JSON PARSING
    # ──────────────────────────────────────────────────────────────

    def start_processing(self):
        if not self.file_paths:
            messagebox.showwarning("Peringatan", "Pilih file terlebih dahulu!")
            return
        if not self.client:
            messagebox.showerror("Error", "API client belum diinisialisasi. Periksa pengaturan API.")
            return
        self.btn_process.state(['disabled'])
        self.progress.start(10)
        threading.Thread(target=self.process_logic, daemon=True).start()

    def update_status(self, msg):
        self.root.after(0, lambda: self.status_var.set(f"● {msg}"))

    def _call_ai_with_retry(self, img, page_num: int, max_retries: int = 3) -> dict:
        """
        Panggil Gemini API dan parse JSON-nya. 
        Retry otomatis hingga max_retries kali jika JSON tidak valid.
        """
        last_error = None

        for attempt in range(1, max_retries + 1):
            if attempt > 1:
                self.update_status(
                    f"Halaman {page_num}: Percobaan ulang {attempt}/{max_retries} "
                    f"(JSON tidak valid sebelumnya)..."
                )

            try:
                response = self.client.models.generate_content(
                    model=self.model_name,
                    contents=[PROMPT_MASTER, img],
                    config={"temperature": 0.1}
                )

                raw = response.text.strip() if response.text else ""
                print(f"[DEBUG] Halaman {page_num}, percobaan {attempt}: "
                      f"Raw response length = {len(raw)}")
                print(f"[DEBUG] 200 karakter pertama: {raw[:200]}")

                # Gunakan parser robust
                page_data = extract_and_fix_json(raw)
                return page_data

            except json.JSONDecodeError as e:
                last_error = e
                print(f"[WARN] Halaman {page_num}, percobaan {attempt}: JSON error — {e}")
                if attempt == max_retries:
                    break
                # Tidak perlu sleep, langsung retry

            except Exception as e:
                # Error API lain (network, quota, dll) — langsung lempar
                raise

        # Semua percobaan gagal
        raise json.JSONDecodeError(
            f"Gagal parse JSON halaman {page_num} setelah {max_retries} percobaan. "
            f"Error terakhir: {last_error}",
            "", 0
        )

    def process_logic(self):
        try:
            from pdf2image import convert_from_path

            self.update_status("Memuat gambar/PDF...")
            all_images = []
            for path in self.file_paths:
                if path.lower().endswith('.pdf'):
                    pages = convert_from_path(path, dpi=200)
                    all_images.extend(pages)
                    self.update_status(f"PDF dikonversi: {len(pages)} halaman")
                else:
                    img = Image.open(path)
                    if img.width < 1000:
                        scale = 1000 / img.width
                        img = img.resize((int(img.width * scale), int(img.height * scale)),
                                         Image.LANCZOS)
                    all_images.append(img)

            self.update_status(
                f"Menganalisis {len(all_images)} halaman dengan AI ({self.model_name})...")

            all_elements = []
            page_layout  = {"orientation": "portrait", "has_letterhead": False}

            for i, img in enumerate(all_images):
                self.update_status(f"Memproses halaman {i+1}/{len(all_images)}...")

                # ← Gunakan fungsi dengan retry & robust JSON parsing
                page_data = self._call_ai_with_retry(img, page_num=i+1, max_retries=3)

                if i == 0:
                    page_layout = page_data.get("page_layout", page_layout)

                elements = page_data.get("elements", [])
                all_elements.extend(elements)

                if i < len(all_images) - 1 and elements:
                    all_elements.append({"type": "page_break"})

            final_data = {"page_layout": page_layout, "elements": all_elements}
            self.update_status("Membuat dokumen Word...")
            self.root.after(0, lambda: self.create_docx(final_data))

        except json.JSONDecodeError as e:
            err_msg = (
                f"AI mengembalikan format JSON tidak valid setelah beberapa percobaan ulang.\n\n"
                f"Saran:\n"
                f"• Coba gunakan model yang lebih canggih (gemini-2.5-pro)\n"
                f"• Periksa kualitas/resolusi gambar\n"
                f"• Pastikan gambar tidak terlalu kompleks\n\n"
                f"File debug disimpan di: debug_last_response.txt\n\n"
                f"Detail teknis: {str(e)[:200]}"
            )
            self.root.after(0, lambda msg=err_msg: messagebox.showerror("Error JSON", msg))
        except Exception as e:
            err_msg = str(e)
            self.root.after(0, lambda msg=err_msg: messagebox.showerror("Error", msg))
        finally:
            self.root.after(0, self.reset_ui)

    # ──────────────────────────────────────────────────────────────
    # DOCX CREATION
    # ──────────────────────────────────────────────────────────────

    def create_docx(self, data):
        doc = Document()

        section  = doc.sections[0]
        layout   = data.get("page_layout", {})
        is_landscape = layout.get("orientation") == "landscape"

        if is_landscape:
            section.orientation  = WD_ORIENT.LANDSCAPE
            section.page_width   = Cm(29.7)
            section.page_height  = Cm(21.0)
            content_width_twip   = PAGE_CONTENT_WIDTH_TWIP_LANDSCAPE
        else:
            section.orientation  = WD_ORIENT.PORTRAIT
            section.page_width   = Cm(21.0)
            section.page_height  = Cm(29.7)
            content_width_twip   = PAGE_CONTENT_WIDTH_TWIP_PORTRAIT

        margin = Inches(1)
        section.left_margin   = margin
        section.right_margin  = margin
        section.top_margin    = margin
        section.bottom_margin = margin

        norm_style = doc.styles['Normal']
        norm_style.font.name = 'Times New Roman'
        norm_style.font.size = Pt(11)

        align_map = {
            "left":    WD_ALIGN_PARAGRAPH.LEFT,
            "center":  WD_ALIGN_PARAGRAPH.CENTER,
            "right":   WD_ALIGN_PARAGRAPH.RIGHT,
            "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
        }

        def apply_run_format(run, el):
            if el.get('bold'):      run.bold      = True
            if el.get('italic'):    run.italic    = True
            if el.get('underline'): run.underline = True
            fs = el.get('font_size', 11)
            if fs: run.font.size = Pt(fs)
            if el.get('is_uppercase'):
                run.text = run.text.upper()

        def apply_para_format(p, el):
            p.alignment = align_map.get(el.get('alignment', 'left'), WD_ALIGN_PARAGRAPH.LEFT)
            fmt = p.paragraph_format
            fmt.space_before = Pt(el.get('space_before', 0))
            fmt.space_after  = Pt(el.get('space_after', 6))
            indent = el.get('indent_level', 0)
            if indent:
                fmt.left_indent = Inches(indent * 0.25)

        for el in data.get("elements", []):
            try:
                e_type  = el.get('type', 'paragraph')
                content = el.get('content', '')

                if e_type == 'page_break':
                    doc.add_page_break()

                elif e_type == 'heading':
                    level = min(max(el.get('level', 1), 1), 6)
                    h = doc.add_heading('', level=level)
                    h.alignment = align_map.get(el.get('alignment', 'center'),
                                                WD_ALIGN_PARAGRAPH.CENTER)
                    run = h.add_run(content.upper() if el.get('is_uppercase') else content)
                    if el.get('font_size'):
                        run.font.size = Pt(el['font_size'])
                    run.bold = True

                elif e_type in ('paragraph', 'field_value'):
                    p = doc.add_paragraph()
                    apply_para_format(p, el)
                    apply_run_format(p.add_run(content), el)

                elif e_type == 'field_label':
                    p   = doc.add_paragraph()
                    apply_para_format(p, el)
                    run = p.add_run(content)
                    run.bold      = True
                    run.font.size = Pt(el.get('font_size', 11))
                    p.paragraph_format.space_after = Pt(2)

                elif e_type == 'signature_block':
                    p = doc.add_paragraph()
                    p.alignment = align_map.get(el.get('alignment', 'right'),
                                                WD_ALIGN_PARAGRAPH.RIGHT)
                    run = p.add_run(content)
                    run.font.size = Pt(el.get('font_size', 11))
                    p.paragraph_format.space_before = Pt(12)

                elif e_type == 'list_item':
                    items      = el.get('items', [content] if content else [])
                    style_name = ('List Bullet' if el.get('list_type') == 'bullet'
                                  else 'List Number')
                    for item in items:
                        li  = doc.add_paragraph(style=style_name)
                        run = li.add_run(item)
                        apply_run_format(run, el)
                        li.alignment = align_map.get(el.get('alignment', 'left'),
                                                     WD_ALIGN_PARAGRAPH.LEFT)

                elif e_type == 'table':
                    build_table(doc, el, content_width_twip)

            except Exception as e:
                print(f"[SKIP] element type='{e_type}' | error: {e}")

        save_path  = self._get_final_save_path()
        doc.save(save_path)
        saved_name = os.path.basename(save_path)
        self._refresh_preview()
        messagebox.showinfo(
            "✅ Sukses",
            f"Dokumen berhasil dibuat!\n\n📄 {saved_name}\n\n📁 {os.path.dirname(save_path)}"
        )
        self.update_status(f"Selesai! → {saved_name}")

    def reset_ui(self):
        self.btn_process.state(['!disabled'])
        self.progress.stop()


if __name__ == "__main__":
    root = tk.Tk()
    app  = OmbudsmanIntelligentDocx(root)
    root.mainloop()