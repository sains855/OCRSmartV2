"""
docx_template.py
Menghasilkan dokumen Word (.docx) dengan format statis formulir Ombudsman RI
dan mengisi data dari hasil OCR.
"""

import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ─── Helpers ─────────────────────────────────────────────────────────────────

def _safe(val, fallback=""):
    if val is None:
        return fallback
    s = str(val).strip()
    return s if s else fallback


def _set_cell_bg(cell, hex_color: str):
    """Set cell background. w:shd must appear after w:tcBorders in tcPr."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    # Insert shd after tcBorders (if present) but before tcMar/vAlign
    insert_pos = len(tcPr)
    for i, child in enumerate(tcPr):
        tag = child.tag.split("}")[-1]
        if tag in ("tcMar", "textDirection", "vAlign"):
            insert_pos = i
            break
    tcPr.insert(insert_pos, shd)


def _set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                       color="000000", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    # Use w:start/w:end (OOXML 2nd ed) in addition to w:left/w:right for compatibility
    sides = [
        ("top", top), ("start", left), ("bottom", bottom),
        ("end", right), ("insideH", False), ("insideV", False),
    ]
    for side, enabled in sides:
        el = OxmlElement(f"w:{side}")
        if enabled:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        borders.append(el)
    # Insert tcBorders after tcW (w:tcW must come first in tcPr)
    # Find position after tcW if exists
    insert_pos = 0
    for i, child in enumerate(tcPr):
        if child.tag.endswith("tcW"):
            insert_pos = i + 1
            break
    tcPr.insert(insert_pos, borders)


def _cell_margins(cell, top=50, bottom=50, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("start", left), ("bottom", bottom), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _para_in_cell(cell, text="", bold=False, italic=False, size=9,
                   align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=0):
    if cell.paragraphs and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_para(doc, text="", bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=0, space_after=2, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _merge_and_set(table, row, col_start, col_end):
    """Merge cells horizontally."""
    cell = table.cell(row, col_start)
    end  = table.cell(row, col_end)
    cell.merge(end)
    return cell


def _add_page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx.oxml.ns.qn("w:lastRenderedPageBreak"))
    # proper page break
    from docx.oxml import OxmlElement
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)


def _page_break(doc):
    from docx.oxml import OxmlElement
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p._p.append(r)
    return p


def _dotted_line(text, doc, size=9):
    """Field label: dots line pattern."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(1)
    r = p.add_run(text)
    r.font.size = Pt(size)
    return p


# ─── Checkbox helper ─────────────────────────────────────────────────────────

def _checkbox(checked: bool) -> str:
    return "☑" if checked else "☐"


# ─── Kop Surat (Header) ──────────────────────────────────────────────────────

def _build_header(doc, logo_left_path=None, logo_right_path=None):
    """Build the official letterhead header."""
    # Header table: logo-left | text center | logo-right
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"

    widths = [2.5, 11.5, 2.5]
    _set_col_widths(tbl, widths)

    row = tbl.rows[0]

    # Remove all borders
    for cell in row.cells:
        _set_cell_borders(cell, False, False, False, False)
        _cell_margins(cell, 0, 0, 80, 80)

    # Left logo
    c_left = row.cells[0]
    c_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if logo_left_path and os.path.exists(logo_left_path):
        p = c_left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_left_path, width=Cm(2.2))
    else:
        _para_in_cell(c_left, "[LOGO KIRI]", size=7,
                       align=WD_ALIGN_PARAGRAPH.CENTER, color="888888")

    # Center text
    c_ctr = row.cells[1]
    c_ctr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _para_in_cell(c_ctr, "OMBUDSMAN REPUBLIK INDONESIA", bold=True, size=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _para_in_cell(c_ctr, "KANTOR PERWAKILAN SULAWESI TENGGARA", bold=True, size=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _para_in_cell(c_ctr,
                  "Jl. Drs. H. Abdullah Silondae No. 114 Kendari, Sulawesi Tenggara 93111,",
                  size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para_in_cell(c_ctr, "Telp/Fax: (0401) 3415554, HP: 08112403737,",
                  size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para_in_cell(c_ctr, "e-mail: sultra@ombudsman.go.id, Website: https://ombudsman.go.id",
                  size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # Right logo
    c_right = row.cells[2]
    c_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if logo_right_path and os.path.exists(logo_right_path):
        p = c_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(logo_right_path, width=Cm(2.0))

    # Horizontal rule below header
    doc.add_paragraph()
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single"); bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)


# ─── Halaman 1 ───────────────────────────────────────────────────────────────

def _build_hal1(doc, data: dict):
    agenda   = data.get("HAL_1_AGENDA",  {}) or {}
    pelapor  = data.get("HAL_1_PELAPOR", {}) or {}
    terlapor = data.get("HAL_1_TERLAPOR",{}) or {}

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("FORMULIR PENERIMAAN LAPORAN/PENGADUAN")
    r.bold = True; r.font.size = Pt(11)
    r.font.underline = True

    # Agenda fields
    def _field_row(label, value):
        """Render label: value on one line."""
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(f"{label}  :  ")
        r1.bold = True; r1.font.size = Pt(9)
        r2 = p.add_run(_safe(value, "………………………"))
        r2.font.size = Pt(9)
        return p

    _field_row("Nomor Agenda ", _safe(agenda.get("nomor_agenda")))
    _field_row("Tanggal Agenda", _safe(agenda.get("tanggal_agenda")))

    doc.add_paragraph()

    # ── IDENTITAS PELAPOR ──
    p = _add_para(doc, "IDENTITAS PELAPOR", bold=True, size=10, space_before=2, space_after=2)

    # Nomor Identitas table
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    widths = [3.0, 4.5, 2.5, 2.5, 3.5]
    _set_col_widths(tbl, widths)
    row = tbl.rows[0]
    for c in row.cells:
        _set_cell_borders(c, False, False, False, False)
        _cell_margins(c, 30, 30, 60, 60)

    _para_in_cell(row.cells[0], "Nomor Identitas", bold=True, size=9)
    _para_in_cell(row.cells[1], f":  {_safe(pelapor.get('nomor_id'))}", size=9)

    jenis = _safe(pelapor.get("jenis_id"), "").upper()
    _para_in_cell(row.cells[2],
                  f"{_checkbox(jenis == 'KTP')} KTP", size=9)
    _para_in_cell(row.cells[3],
                  f"{_checkbox('KITAS' in jenis or 'KITAP' in jenis)} KITAS/KITAP", size=9)
    _para_in_cell(row.cells[4],
                  f"☐ Lainnya ………………", size=9)

    # Personal fields
    fields_pelapor = [
        ("Nama Lengkap",           pelapor.get("nama_lengkap")),
        ("Tempat, Tgl Lahir",      pelapor.get("tempat_tgl_lahir")),
        ("Alamat Korespondensi",   pelapor.get("alamat_korespondensi")),
    ]
    for label, val in fields_pelapor:
        tbl2 = doc.add_table(rows=1, cols=2)
        tbl2.style = "Table Grid"
        _set_col_widths(tbl2, [3.5, 12.5])
        r2 = tbl2.rows[0]
        for c in r2.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 25, 25, 60, 60)
        _para_in_cell(r2.cells[0], label, bold=True, size=9)
        _para_in_cell(r2.cells[1], f":  {_safe(val)}", size=9)

    # Status perkawinan
    nikah = _safe(pelapor.get("status_perkawinan"), "").lower()
    tbl3 = doc.add_table(rows=1, cols=2)
    tbl3.style = "Table Grid"
    _set_col_widths(tbl3, [3.5, 12.5])
    r3 = tbl3.rows[0]
    for c in r3.cells:
        _set_cell_borders(c, False, False, False, False)
        _cell_margins(c, 25, 25, 60, 60)
    _para_in_cell(r3.cells[0], "Status Perkawinan", bold=True, size=9)

    status_txt = (
        f": {_checkbox('kawin' in nikah and 'belum' not in nikah and 'cerai' not in nikah)} Kawin   "
        f"{_checkbox('belum' in nikah)} Belum Kawin   "
        f"{_checkbox('mati' in nikah)} Cerai Mati   "
        f"{_checkbox('hidup' in nikah)} Cerai Hidup"
    )
    _para_in_cell(r3.cells[1], status_txt, size=9)

    # Pekerjaan & Pendidikan
    for label, val in [("Pekerjaan", pelapor.get("pekerjaan")),
                        ("Pendidikan Terakhir", pelapor.get("pendidikan_terakhir")),
                        ("Nomor Telepon Aktif", pelapor.get("telp")),
                        ("Email Aktif",         pelapor.get("email"))]:
        tbl4 = doc.add_table(rows=1, cols=2)
        tbl4.style = "Table Grid"
        _set_col_widths(tbl4, [3.5, 12.5])
        r4 = tbl4.rows[0]
        for c in r4.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 25, 25, 60, 60)
        _para_in_cell(r4.cells[0], label, bold=True, size=9)
        _para_in_cell(r4.cells[1], f":  {_safe(val)}", size=9)

    doc.add_paragraph()

    # ── TERLAPOR ──
    _add_para(doc, "TERLAPOR", bold=True, size=10, space_before=2, space_after=2)

    for label, val in [
        ("Nama Terlapor",    terlapor.get("nama_terlapor")),
        ("Jabatan Terlapor", terlapor.get("jabatan_terlapor")),
        ("Instansi Terlapor",terlapor.get("instansi_terlapor")),
        ("Alamat Terlapor",  terlapor.get("alamat_terlapor")),
    ]:
        tbl5 = doc.add_table(rows=1, cols=2)
        tbl5.style = "Table Grid"
        _set_col_widths(tbl5, [3.5, 12.5])
        r5 = tbl5.rows[0]
        for c in r5.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 25, 25, 60, 60)
        _para_in_cell(r5.cells[0], label, bold=True, size=9)
        _para_in_cell(r5.cells[1], f":  {_safe(val)}", size=9)

    doc.add_paragraph()

    # ── WAKTU PERISTIWA ──
    kronologi = data.get("HAL_2_3_KRONOLOGI", []) or []
    first_date = ""
    if isinstance(kronologi, list) and kronologi:
        first_date = _safe(kronologi[0].get("tanggal"))
    elif isinstance(kronologi, dict):
        first_date = _safe(kronologi.get("waktu_peristiwa"))

    tbl6 = doc.add_table(rows=1, cols=2)
    tbl6.style = "Table Grid"
    _set_col_widths(tbl6, [3.8, 12.2])
    r6 = tbl6.rows[0]
    for c in r6.cells:
        _set_cell_borders(c, False, False, False, False)
        _cell_margins(c, 25, 25, 60, 60)
    _para_in_cell(r6.cells[0], "WAKTU PERISTIWA", bold=True, size=9)
    _para_in_cell(r6.cells[1],
        f"Kapan peristiwa, tindakan atau keputusan yang dilaporkan terjadi?\n"
        f"Tanggal, Bulan, Tahun  :  {first_date}", size=9)

    # Footer page 1
    _build_page_footer(doc, 1)


# ─── Halaman 2 ───────────────────────────────────────────────────────────────

def _build_hal2(doc, data: dict):
    validasi  = data.get("HAL_2_VALIDASI",      {}) or {}
    kronologi = data.get("HAL_2_3_KRONOLOGI",   []) or []

    # ── Validasi table ──
    doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [5.5, 1.2, 1.5, 7.8])

    border = {"color": "000000", "size": 4}

    # Header row
    hdr_cells = tbl.rows[0].cells
    for c in hdr_cells:
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 60, 100, 100)
        _set_cell_bg(c, "D9D9D9")
    _para_in_cell(hdr_cells[0], "Pertanyaan", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr_cells[1], "Ya",  bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr_cells[2], "Tidak", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr_cells[3], "Keterangan", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Row 1 – sudah lapor instansi
    sudah_lapor = _safe(validasi.get("sudah_lapor_instansi_terkait"), "Tidak").lower() == "ya"
    r1 = tbl.rows[1]
    for c in r1.cells:
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 60, 100, 100)
    _para_in_cell(r1.cells[0],
        "Apakah saudara sudah menyampaikan Laporan kepada instansi Terlapor?", size=9)
    _para_in_cell(r1.cells[1], _checkbox(sudah_lapor), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(r1.cells[2], _checkbox(not sudah_lapor), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Keterangan (instansi)
    instansi   = _safe(validasi.get("detail_instansi"))
    tgl_laporan= _safe(validasi.get("detail_tgl_laporan"))
    melalui    = _safe(validasi.get("detail_melalui"), "").lower()
    ket_text   = (
        f"Instansi    : {instansi}\n"
        f"Tgl/bln/thn : {tgl_laporan}\n"
        f"Melalui     : "
        f"{_checkbox('surat' in melalui)} Surat  "
        f"{_checkbox('langsung' in melalui)} Datang Langsung  "
        f"{_checkbox('telepon' in melalui)} Telepon  "
        f"☐ Lainnya"
    )
    _para_in_cell(r1.cells[3], ket_text, size=8)

    # Row 2 – gugatan pengadilan
    digugat = _safe(validasi.get("substansi_digugat_pengadilan"), "Tidak").lower() == "ya"
    r2 = tbl.rows[2]
    for c in r2.cells:
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 60, 100, 100)
    _para_in_cell(r2.cells[0],
        "Apakah substansi laporan sudah pernah atau sedang diajukan gugatan ke pengadilan?", size=9)
    _para_in_cell(r2.cells[1], _checkbox(digugat),     size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(r2.cells[2], _checkbox(not digugat), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

    peng  = _safe(validasi.get("nama_pengadilan"))
    noreg = _safe(validasi.get("nomor_register_perkara"))
    _para_in_cell(r2.cells[3],
        f"Pengadilan           : {peng}\nNomor Register Perkara : {noreg}", size=8)

    doc.add_paragraph()

    # ── Kronologi table (hal 2-3) ──
    _build_kronologi_table(doc, kronologi)

    _build_page_footer(doc, 2)


# ─── Kronologi table ─────────────────────────────────────────────────────────

def _build_kronologi_table(doc, kronologi):
    """Build the event chronology table (spans hal 2-3)."""
    _add_para(doc, "URAIAN PERISTIWA (KRONOLOGI)", bold=True, size=10,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)

    # Normalize kronologi
    rows_data = []
    if isinstance(kronologi, list):
        for item in kronologi:
            if isinstance(item, dict):
                rows_data.append((
                    _safe(item.get("tanggal")),
                    _safe(item.get("peristiwa")),
                    _safe(item.get("catatan_bukti")),
                ))
    elif isinstance(kronologi, dict):
        rows_data.append((
            _safe(kronologi.get("waktu_peristiwa")),
            _safe(kronologi.get("uraian_peristiwa_lengkap")),
            "",
        ))

    # Minimum 6 rows
    while len(rows_data) < 6:
        rows_data.append(("", "", ""))

    # Table: header + data rows
    tbl = doc.add_table(rows=1 + len(rows_data), cols=3)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [3.0, 10.0, 3.5])

    # Header
    hdr = tbl.rows[0]
    _set_cell_bg(hdr.cells[0], "D9D9D9")
    _set_cell_bg(hdr.cells[1], "D9D9D9")
    _set_cell_bg(hdr.cells[2], "D9D9D9")
    for c in hdr.cells:
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 30, 100, 100)
    _para_in_cell(hdr.cells[0], "Tanggal\n(hh/bb/tttt)", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr.cells[1], "Peristiwa\n(5W 1 H)", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr.cells[2], "Catatan/Bukti", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (tgl, peristiwa, catatan) in enumerate(rows_data):
        row = tbl.rows[i + 1]
        for c in row.cells:
            _set_cell_borders(c, True, True, True, True)
            _cell_margins(c, 40, 40, 100, 100)
        # Set minimum row height
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "800")  # ~0.56cm minimum
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)

        _para_in_cell(row.cells[0], tgl,      size=9)
        _para_in_cell(row.cells[1], peristiwa, size=9)
        _para_in_cell(row.cells[2], catatan,   size=9)


# ─── Halaman 3 (continuation) ────────────────────────────────────────────────

def _build_hal3_footer(doc):
    _build_page_footer(doc, 3)


# ─── Halaman 4 ───────────────────────────────────────────────────────────────

def _build_hal4(doc, data: dict):
    akhir = data.get("HAL_4_AKHIR", {}) or {}
    ttd   = data.get("HAL_4_TTD",   {}) or {}

    # Harapan Pelapor
    doc.add_paragraph()
    _add_para(doc, "HARAPAN PELAPOR:", bold=True, size=10, space_before=2, space_after=2)

    harapan = _safe(akhir.get("harapan_pelapor"), "………………………………………………………………………")
    tbl1 = doc.add_table(rows=1, cols=1)
    tbl1.style = "Table Grid"
    _set_col_widths(tbl1, [16.5])
    c = tbl1.rows[0].cells[0]
    _set_cell_borders(c, True, True, True, True)
    _cell_margins(c, 80, 80, 120, 120)
    # Set height
    trPr = tbl1.rows[0]._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), "1200"); trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)
    _para_in_cell(c, harapan, size=9)

    doc.add_paragraph()

    # Dokumen pendukung + identitas dirahasiakan
    tbl2 = doc.add_table(rows=2, cols=4)
    tbl2.style = "Table Grid"
    _set_col_widths(tbl2, [7.0, 2.5, 2.5, 4.5])

    dok = _safe(akhir.get("dokumen_pendukung"), "").lower()
    ada_dok  = "ada" in dok and "tidak" not in dok
    rahasia  = _safe(akhir.get("permintaan_identitas_dirahasiakan"), "").lower() == "ya"

    # Remove all outer borders, keep inner
    for row in tbl2.rows:
        for c in row.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 40, 40, 80, 80)

    r0 = tbl2.rows[0]
    _para_in_cell(r0.cells[0],
        "DOKUMEN PENDUKUNG:\nIdentitas Pelapor (KTP / KITAP / KITAS / Lainnya)", bold=True, size=9)
    _para_in_cell(r0.cells[1], _checkbox(ada_dok) + " Ada",       size=10)
    _para_in_cell(r0.cells[2], _checkbox(not ada_dok) + " Tidak Ada", size=10)
    _para_in_cell(r0.cells[3], "", size=9)

    r1 = tbl2.rows[1]
    _para_in_cell(r1.cells[0], "Permintaan Identitas Dirahasiakan:", bold=True, size=9)
    _para_in_cell(r1.cells[1], _checkbox(rahasia) + " Ya",       size=10)
    _para_in_cell(r1.cells[2], _checkbox(not rahasia) + " Tidak", size=10)
    _para_in_cell(r1.cells[3], "", size=9)

    doc.add_paragraph()

    # Catatan
    _add_para(doc, "CATATAN:", bold=True, size=10, space_before=2, space_after=2)
    catatan = _safe(akhir.get("catatan_tambahan"), "")
    tbl3 = doc.add_table(rows=1, cols=1)
    tbl3.style = "Table Grid"
    _set_col_widths(tbl3, [16.5])
    c3 = tbl3.rows[0].cells[0]
    _set_cell_borders(c3, True, True, True, True)
    _cell_margins(c3, 80, 80, 120, 120)
    trPr3 = tbl3.rows[0]._tr.get_or_add_trPr()
    trH3 = OxmlElement("w:trHeight")
    trH3.set(qn("w:val"), "1000"); trH3.set(qn("w:hRule"), "atLeast")
    trPr3.append(trH3)
    _para_in_cell(c3, catatan, size=9)

    doc.add_paragraph()

    # Tanda tangan
    kota_tgl  = _safe(ttd.get("kota_tanggal_ttd"), "Kendari, ………………………… 20 …..")
    nm_pelapor = _safe(ttd.get("nama_pelapor_ttd"))
    nm_petugas = _safe(ttd.get("nama_penerima_laporan_petugas"))

    tbl4 = doc.add_table(rows=4, cols=3)
    tbl4.style = "Table Grid"
    _set_col_widths(tbl4, [5.5, 6.0, 5.0])
    for row in tbl4.rows:
        for c in row.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 40, 40, 80, 80)

    _para_in_cell(tbl4.rows[0].cells[0], "", size=9)
    _para_in_cell(tbl4.rows[0].cells[1], "", size=9)
    _para_in_cell(tbl4.rows[0].cells[2], kota_tgl, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    _para_in_cell(tbl4.rows[1].cells[0], "Pelapor", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(tbl4.rows[1].cells[1], "", size=9)
    _para_in_cell(tbl4.rows[1].cells[2], "Penerima Laporan", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Signature space rows
    for row in [tbl4.rows[2]]:
        trPr = row._tr.get_or_add_trPr()
        trH = OxmlElement("w:trHeight")
        trH.set(qn("w:val"), "1200"); trH.set(qn("w:hRule"), "atLeast")
        trPr.append(trH)
        for c in row.cells:
            _para_in_cell(c, "", size=9)

    _para_in_cell(tbl4.rows[3].cells[0], nm_pelapor, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(tbl4.rows[3].cells[1], "", size=9)
    _para_in_cell(tbl4.rows[3].cells[2], nm_petugas, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _add_para(doc, "Catatan: Dokumen yang telah diserahkan menjadi milik Ombudsman RI",
               size=8, space_before=2, space_after=2)

    _build_page_footer(doc, 4)


# ─── Page footer ─────────────────────────────────────────────────────────────

def _build_page_footer(doc, page_num: int):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single"); top.set(qn("w:sz"), "6")
    top.set(qn("w:color"), "000000")
    pBdr.append(top); pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)

    txt = (
        f"{page_num} | Unit Penerimaan dan Verifikasi Laporan (PVL)\n"
        "OMBUDSMAN RI PERWAKILAN PROVINSI SULTRA"
    )
    run = p.add_run(txt)
    run.font.size = Pt(8)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ─── Main builder ────────────────────────────────────────────────────────────

def create_ombudsman_docx(forms: list, output_path: str):
    """
    Create a Word document with one full form (4 pages) per entry in `forms`.
    """
    # Locate logos relative to this file
    here         = os.path.dirname(os.path.abspath(__file__))
    logo_left    = os.path.join(here, "logo-kiri1.jpg")

    doc = Document()

    # ── Page setup (A4) ──
    from docx.oxml import OxmlElement
    section = doc.sections[0]
    section.page_width  = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    for form_idx, form in enumerate(forms):
        if form_idx > 0:
            _page_break(doc)

        # ── Halaman 1 ──
        _build_header(doc, logo_left)
        _build_hal1(doc, form)

        _page_break(doc)

        # ── Halaman 2 ──
        _build_hal2(doc, form)

        _page_break(doc)

        # ── Halaman 3 – continuation of kronologi (blank continuation) ──
        _add_para(doc, "URAIAN PERISTIWA (KRONOLOGI) – Lanjutan", bold=True,
                   size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4)

        # Extra empty kronologi table for continuation pages
        empty_kron = [("", "", "") for _ in range(8)]
        tbl = doc.add_table(rows=1 + len(empty_kron), cols=3)
        tbl.style = "Table Grid"
        _set_col_widths(tbl, [3.0, 10.0, 3.5])
        hdr = tbl.rows[0]
        for c in hdr.cells:
            _set_cell_bg(c, "D9D9D9")
            _set_cell_borders(c, True, True, True, True)
            _cell_margins(c, 60, 30, 100, 100)
        _para_in_cell(hdr.cells[0], "Tanggal\n(hh/bb/tttt)", bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _para_in_cell(hdr.cells[1], "Peristiwa\n(5W 1 H)", bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _para_in_cell(hdr.cells[2], "Catatan/Bukti", bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for i in range(len(empty_kron)):
            row = tbl.rows[i + 1]
            for c in row.cells:
                _set_cell_borders(c, True, True, True, True)
                _cell_margins(c, 40, 40, 100, 100)
            trPr = row._tr.get_or_add_trPr()
            trH = OxmlElement("w:trHeight")
            trH.set(qn("w:val"), "800"); trH.set(qn("w:hRule"), "atLeast")
            trPr.append(trH)
            for c in row.cells:
                _para_in_cell(c, "", size=9)

        _build_page_footer(doc, 3)
        _page_break(doc)

        # ── Halaman 4 ──
        _build_hal4(doc, form)

    doc.save(output_path)
    # Fix w:zoom setting (python-docx generates invalid zoom without percent attr)
    import zipfile as _zf, shutil as _sh
    from lxml import etree
    tmp = output_path + ".tmp"
    _sh.copy(output_path, tmp)
    with _zf.ZipFile(tmp, "r") as zin, _zf.ZipFile(output_path, "w", _zf.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                tree = etree.fromstring(data)
                ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                for z in tree.findall(f"{{{ns}}}zoom"):
                    tree.remove(z)
                data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    import os as _os; _os.remove(tmp)
    return output_path