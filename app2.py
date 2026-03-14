"""
ombudsman_app.py
Aplikasi tunggal: GUI Tkinter + Generator Word (.docx) untuk Formulir Ombudsman RI
Gabungan dari app.py dan docx_template.py

PERBAIKAN: Mendukung banyak formulir dalam satu PDF
- 1 formulir = 4 halaman (hal 1–4)
- Jika PDF berisi 8 halaman → 2 formulir, 12 halaman → 3 formulir, dst.
- Setiap set 4 halaman dikirim ke AI secara terpisah → hasil digabung

PERBAIKAN v2.1: Kompatibel dengan PyInstaller (--onefile & --onedir)
- Semua aset (logo, config) diakses via resource_path()
- Logo dibundel sebagai --add-data di PyInstaller spec
"""

import os
import sys
import re
import json
import threading
import zipfile as _zf
import shutil as _sh
import tkinter as tk
from tkinter import filedialog, messagebox, ttk
from math import ceil

from PIL import Image, ImageTk
from pdf2image import convert_from_path
from google import genai
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ═══════════════════════════════════════════════════════════════════════════════
# PYINSTALLER-COMPATIBLE RESOURCE PATH
# ═══════════════════════════════════════════════════════════════════════════════

def resource_path(relative_path: str) -> str:
    """
    Dapatkan path absolut ke resource.

    - Saat berjalan normal (development): gunakan direktori script.
    - Saat di-bundle PyInstaller (--onefile / --onedir): gunakan sys._MEIPASS
      yang merupakan folder sementara tempat PyInstaller mengekstrak aset.

    Cara menambahkan aset di PyInstaller:
        pyinstaller --add-data "logo-kiri1.jpg;." --add-data "logo-kanan.png;." ombudsman_app.py

    Atau di file .spec:
        datas=[
            ('logo-kiri1.jpg', '.'),
            ('logo-kanan.png', '.'),
        ]
    """
    if getattr(sys, 'frozen', False):
        # Aplikasi sedang berjalan sebagai bundle PyInstaller
        base_dir = sys._MEIPASS
    else:
        # Aplikasi berjalan langsung sebagai script Python
        base_dir = os.path.dirname(os.path.abspath(__file__))

    return os.path.join(base_dir, relative_path)


def config_path() -> str:
    """
    Path untuk config.json — SELALU di direktori yang bisa ditulis
    (bukan _MEIPASS yang read-only saat --onefile).

    Urutan prioritas:
    1. Direktori yang sama dengan .exe (untuk distribusi portable)
    2. Direktori script (untuk development)
    """
    if getattr(sys, 'frozen', False):
        # Saat di-bundle: simpan config di sebelah file .exe
        exe_dir = os.path.dirname(sys.executable)
        return os.path.join(exe_dir, "config.json")
    else:
        return os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.json")


# ═══════════════════════════════════════════════════════════════════════════════
# KONFIGURASI & KONSTANTA
# ═══════════════════════════════════════════════════════════════════════════════

CONFIG_FILE = config_path()

# Jumlah halaman per satu formulir Ombudsman
PAGES_PER_FORM = 4

EXTRACTION_PROMPT = """Kamu adalah sistem OCR presisi tinggi untuk formulir Ombudsman Republik Indonesia.
Kamu menerima TEPAT 4 halaman yang merupakan SATU formulir lengkap.
Baca SETIAP tulisan dengan sangat teliti lalu hasilkan JSON bersih.

ATURAN:
1. Tulisan tangan: baca apa adanya, jangan perbaiki ejaan.
2. Checkbox: tulis HANYA opsi yang dicentang (✓). Jika tidak ada, isi "-".
3. Kata dicoret: ABAIKAN. Ambil hanya kata yang BERSIH/tidak dicoret.
4. Field kosong: isi dengan null.
5. Nomor: salin digit per digit.

FORMAT OUTPUT (wajib JSON murni, tanpa markdown):
{
  "HAL_1_AGENDA": {
    "nomor_agenda": "...",
    "tanggal_agenda": "..."
  },
  "HAL_1_PELAPOR": {
    "jenis_id": "KTP atau KITAS",
    "nomor_id": "...",
    "nama_lengkap": "...",
    "tempat_tgl_lahir": "...",
    "alamat_korespondensi": "...",
    "status_perkawinan": "Kawin / Belum Kawin / Cerai Mati / Cerai Hidup",
    "pekerjaan": "...",
    "pendidikan_terakhir": "...",
    "telp": "...",
    "email": "..."
  },
  "HAL_1_TERLAPOR": {
    "nama_terlapor": "...",
    "jabatan_terlapor": "...",
    "instansi_terlapor": "...",
    "alamat_terlapor": "..."
  },
  "HAL_2_VALIDASI": {
    "sudah_lapor_instansi_terkait": "Ya atau Tidak",
    "detail_instansi": "nama instansi",
    "detail_tgl_laporan": "tanggal laporan",
    "detail_melalui": "Surat / Datang Langsung / Telepon / Lainnya",
    "substansi_digugat_pengadilan": "Ya atau Tidak",
    "nama_pengadilan": "...",
    "nomor_register_perkara": "..."
  },
  "HAL_2_3_KRONOLOGI": [
    {
      "tanggal": "...",
      "peristiwa": "...",
      "catatan_bukti": "..."
    }
  ],
  "HAL_4_AKHIR": {
    "harapan_pelapor": "...",
    "dokumen_pendukung": "Ada atau Tidak Ada",
    "permintaan_identitas_dirahasiakan": "Ya atau Tidak",
    "catatan_tambahan": "..."
  },
  "HAL_4_TTD": {
    "kota_tanggal_ttd": "...",
    "nama_pelapor_ttd": "...",
    "nama_penerima_laporan_petugas": "..."
  }
}

KELUARKAN HANYA JSON MURNI. Jangan tambahkan markdown, komentar, atau teks apapun di luar JSON.
"""

# Color palette
C = {
    "bg":             "#F8FAFC",
    "surface":        "#FFFFFF",
    "surface2":       "#F1F5F9",
    "border":         "#CBD5E1",
    "accent":         "#2563EB",
    "accent_dark":    "#1D4ED8",
    "success":        "#059669",
    "warning":        "#D97706",
    "danger":         "#DC2626",
    "text_primary":   "#0F172A",
    "text_secondary": "#475569",
    "text_muted":     "#94A3B8",
    "white":          "#FFFFFF",
}

FONT_BODY  = ("Segoe UI", 9)
FONT_SMALL = ("Segoe UI", 8)
FONT_BTN   = ("Segoe UI", 9, "bold")

# ── Path logo menggunakan resource_path() agar kompatibel dengan PyInstaller ──
LOGO_LEFT_PATH  = resource_path("logo-kiri1.jpg")
LOGO_RIGHT_PATH = resource_path("logo-kanan.png")
LOGO_SIZE       = (56, 56)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX HELPERS
# ═══════════════════════════════════════════════════════════════════════════════

def _safe(val, fallback=""):
    if val is None:
        return fallback
    s = str(val).strip()
    return s if s else fallback


def _set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    insert_pos = len(tcPr)
    for i, child in enumerate(tcPr):
        tag = child.tag.split("}")[-1]
        if tag in ("tcMar", "textDirection", "vAlign"):
            insert_pos = i
            break
    tcPr.insert(insert_pos, shd)


def _set_cell_borders(cell, top=True, bottom=True, left=True, right=True,
                       color="000000", size=4):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    sides = [
        ("top", top), ("start", left), ("bottom", bottom),
        ("end", right), ("insideH", False), ("insideV", False),
    ]
    for side, enabled in sides:
        el = OxmlElement(f"w:{side}")
        if enabled:
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(size))
            el.set(qn("w:color"), color)
        else:
            el.set(qn("w:val"), "none")
        borders.append(el)
    insert_pos = 0
    for i, child in enumerate(tcPr):
        if child.tag.endswith("tcW"):
            insert_pos = i + 1
            break
    tcPr.insert(insert_pos, borders)


def _cell_margins(cell, top=50, bottom=50, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("start", left), ("bottom", bottom), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        mar.append(el)
    tcPr.append(mar)


def _para_in_cell(cell, text="", bold=False, italic=False, size=9,
                   align=WD_ALIGN_PARAGRAPH.LEFT, color=None, space_before=0, space_after=0):
    if cell.paragraphs and cell.paragraphs[0].text == "":
        p = cell.paragraphs[0]
    else:
        p = cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_para(doc, text="", bold=False, size=9, align=WD_ALIGN_PARAGRAPH.LEFT,
               space_before=0, space_after=2, color=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths_cm):
                cell.width = Cm(widths_cm[i])


def _page_break(doc):
    p = doc.add_paragraph()
    r = OxmlElement("w:r")
    br = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    r.append(br)
    p._p.append(r)
    return p


def _checkbox(checked: bool) -> str:
    return "☑" if checked else "☐"


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER — KOP SURAT
# ═══════════════════════════════════════════════════════════════════════════════

def _build_header(doc, logo_left_path=None, logo_right_path=None):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [2.5, 11.5, 2.5])

    row = tbl.rows[0]
    for cell in row.cells:
        _set_cell_borders(cell, False, False, False, False)
        _cell_margins(cell, 0, 0, 80, 80)

    # Kiri
    c_left = row.cells[0]
    c_left.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if logo_left_path and os.path.exists(logo_left_path):
        p = c_left.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo_left_path, width=Cm(2.2))
    else:
        _para_in_cell(c_left, "[LOGO KIRI]", size=7,
                       align=WD_ALIGN_PARAGRAPH.CENTER, color="888888")

    # Tengah
    c_ctr = row.cells[1]
    c_ctr.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    _para_in_cell(c_ctr, "OMBUDSMAN REPUBLIK INDONESIA", bold=True, size=12,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _para_in_cell(c_ctr, "KANTOR PERWAKILAN SULAWESI TENGGARA", bold=True, size=11,
                   align=WD_ALIGN_PARAGRAPH.CENTER, space_after=1)
    _para_in_cell(c_ctr,
                  "Jl. Drs. H. Abdullah Silondae No. 114 Kendari, Sulawesi Tenggara 93111,",
                  size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para_in_cell(c_ctr, "Telp/Fax: (0401) 3415554, HP: 08112403737,",
                  size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0)
    _para_in_cell(c_ctr, "e-mail: sultra@ombudsman.go.id, Website: https://ombudsman.go.id",
                  size=8, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)

    # Kanan
    c_right = row.cells[2]
    c_right.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if logo_right_path and os.path.exists(logo_right_path):
        p = c_right.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(logo_right_path, width=Cm(2.0))

    # Garis bawah header
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:color"), "000000")
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER — HALAMAN 1
# ═══════════════════════════════════════════════════════════════════════════════

def _build_hal1(doc, data: dict):
    agenda   = data.get("HAL_1_AGENDA",  {}) or {}
    pelapor  = data.get("HAL_1_PELAPOR", {}) or {}
    terlapor = data.get("HAL_1_TERLAPOR", {}) or {}

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run("FORMULIR PENERIMAAN LAPORAN/PENGADUAN")
    r.bold = True
    r.font.size = Pt(11)
    r.font.underline = True

    def _field_row(label, value):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(f"{label}  :  ")
        r1.bold = True; r1.font.size = Pt(9)
        r2 = p.add_run(_safe(value, "………………………"))
        r2.font.size = Pt(9)

    _field_row("Nomor Agenda ", _safe(agenda.get("nomor_agenda")))
    _field_row("Tanggal Agenda", _safe(agenda.get("tanggal_agenda")))
    doc.add_paragraph()

    _add_para(doc, "IDENTITAS PELAPOR", bold=True, size=10, space_before=2, space_after=2)

    # Nomor Identitas
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [3.0, 4.5, 2.5, 2.5, 3.5])
    row = tbl.rows[0]
    for c in row.cells:
        _set_cell_borders(c, False, False, False, False)
        _cell_margins(c, 30, 30, 60, 60)
    _para_in_cell(row.cells[0], "Nomor Identitas", bold=True, size=9)
    _para_in_cell(row.cells[1], f":  {_safe(pelapor.get('nomor_id'))}", size=9)
    jenis = _safe(pelapor.get("jenis_id"), "").upper()
    _para_in_cell(row.cells[2], f"{_checkbox(jenis == 'KTP')} KTP", size=9)
    _para_in_cell(row.cells[3], f"{_checkbox('KITAS' in jenis or 'KITAP' in jenis)} KITAS/KITAP", size=9)
    _para_in_cell(row.cells[4], "☐ Lainnya ………………", size=9)

    # Field-field pelapor
    for label, val in [
        ("Nama Lengkap",         pelapor.get("nama_lengkap")),
        ("Tempat, Tgl Lahir",    pelapor.get("tempat_tgl_lahir")),
        ("Alamat Korespondensi", pelapor.get("alamat_korespondensi")),
    ]:
        t = doc.add_table(rows=1, cols=2)
        t.style = "Table Grid"
        _set_col_widths(t, [3.5, 12.5])
        r2 = t.rows[0]
        for c in r2.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 25, 25, 60, 60)
        _para_in_cell(r2.cells[0], label, bold=True, size=9)
        _para_in_cell(r2.cells[1], f":  {_safe(val)}", size=9)

    # Status perkawinan
    nikah = _safe(pelapor.get("status_perkawinan"), "").lower()
    t3 = doc.add_table(rows=1, cols=2)
    t3.style = "Table Grid"
    _set_col_widths(t3, [3.5, 12.5])
    r3 = t3.rows[0]
    for c in r3.cells:
        _set_cell_borders(c, False, False, False, False)
        _cell_margins(c, 25, 25, 60, 60)
    _para_in_cell(r3.cells[0], "Status Perkawinan", bold=True, size=9)
    status_txt = (
        f": {_checkbox('kawin' in nikah and 'belum' not in nikah and 'cerai' not in nikah)} Kawin   "
        f"{_checkbox('belum' in nikah)} Belum Kawin   "
        f"{_checkbox('mati' in nikah)} Cerai Mati   "
        f"{_checkbox('hidup' in nikah)} Cerai Hidup"
    )
    _para_in_cell(r3.cells[1], status_txt, size=9)

    # Pekerjaan, pendidikan, telepon, email
    for label, val in [
        ("Pekerjaan",            pelapor.get("pekerjaan")),
        ("Pendidikan Terakhir",  pelapor.get("pendidikan_terakhir")),
        ("Nomor Telepon Aktif",  pelapor.get("telp")),
        ("Email Aktif",          pelapor.get("email")),
    ]:
        t4 = doc.add_table(rows=1, cols=2)
        t4.style = "Table Grid"
        _set_col_widths(t4, [3.5, 12.5])
        r4 = t4.rows[0]
        for c in r4.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 25, 25, 60, 60)
        _para_in_cell(r4.cells[0], label, bold=True, size=9)
        _para_in_cell(r4.cells[1], f":  {_safe(val)}", size=9)

    doc.add_paragraph()
    _add_para(doc, "TERLAPOR", bold=True, size=10, space_before=2, space_after=2)

    for label, val in [
        ("Nama Terlapor",     terlapor.get("nama_terlapor")),
        ("Jabatan Terlapor",  terlapor.get("jabatan_terlapor")),
        ("Instansi Terlapor", terlapor.get("instansi_terlapor")),
        ("Alamat Terlapor",   terlapor.get("alamat_terlapor")),
    ]:
        t5 = doc.add_table(rows=1, cols=2)
        t5.style = "Table Grid"
        _set_col_widths(t5, [3.5, 12.5])
        r5 = t5.rows[0]
        for c in r5.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 25, 25, 60, 60)
        _para_in_cell(r5.cells[0], label, bold=True, size=9)
        _para_in_cell(r5.cells[1], f":  {_safe(val)}", size=9)

    doc.add_paragraph()

    # Waktu peristiwa
    kronologi = data.get("HAL_2_3_KRONOLOGI", []) or []
    first_date = ""
    if isinstance(kronologi, list) and kronologi:
        first_date = _safe(kronologi[0].get("tanggal"))
    elif isinstance(kronologi, dict):
        first_date = _safe(kronologi.get("waktu_peristiwa"))

    t6 = doc.add_table(rows=1, cols=2)
    t6.style = "Table Grid"
    _set_col_widths(t6, [3.8, 12.2])
    r6 = t6.rows[0]
    for c in r6.cells:
        _set_cell_borders(c, False, False, False, False)
        _cell_margins(c, 25, 25, 60, 60)
    _para_in_cell(r6.cells[0], "WAKTU PERISTIWA", bold=True, size=9)
    _para_in_cell(r6.cells[1],
        f"Kapan peristiwa, tindakan atau keputusan yang dilaporkan terjadi?\n"
        f"Tanggal, Bulan, Tahun  :  {first_date}", size=9)

    _build_page_footer(doc, 1)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER — HALAMAN 2
# ═══════════════════════════════════════════════════════════════════════════════

def _build_hal2(doc, data: dict):
    validasi  = data.get("HAL_2_VALIDASI",    {}) or {}
    kronologi = data.get("HAL_2_3_KRONOLOGI", []) or []

    doc.add_paragraph()
    tbl = doc.add_table(rows=3, cols=4)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [5.5, 1.2, 1.5, 7.8])

    # Header
    hdr_cells = tbl.rows[0].cells
    for c in hdr_cells:
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 60, 100, 100)
        _set_cell_bg(c, "D9D9D9")
    _para_in_cell(hdr_cells[0], "Pertanyaan", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr_cells[1], "Ya",  bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr_cells[2], "Tidak", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr_cells[3], "Keterangan", bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    # Baris 1 – sudah lapor instansi
    sudah_lapor = _safe(validasi.get("sudah_lapor_instansi_terkait"), "Tidak").lower() == "ya"
    r1 = tbl.rows[1]
    for c in r1.cells:
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 60, 100, 100)
    _para_in_cell(r1.cells[0],
        "Apakah saudara sudah menyampaikan Laporan kepada instansi Terlapor?", size=9)
    _para_in_cell(r1.cells[1], _checkbox(sudah_lapor),     size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(r1.cells[2], _checkbox(not sudah_lapor), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    instansi    = _safe(validasi.get("detail_instansi"))
    tgl_laporan = _safe(validasi.get("detail_tgl_laporan"))
    melalui     = _safe(validasi.get("detail_melalui"), "").lower()
    ket_text = (
        f"Instansi    : {instansi}\n"
        f"Tgl/bln/thn : {tgl_laporan}\n"
        f"Melalui     : "
        f"{_checkbox('surat' in melalui)} Surat  "
        f"{_checkbox('langsung' in melalui)} Datang Langsung  "
        f"{_checkbox('telepon' in melalui)} Telepon  "
        f"☐ Lainnya"
    )
    _para_in_cell(r1.cells[3], ket_text, size=8)

    # Baris 2 – gugatan pengadilan
    digugat = _safe(validasi.get("substansi_digugat_pengadilan"), "Tidak").lower() == "ya"
    r2 = tbl.rows[2]
    for c in r2.cells:
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 60, 100, 100)
    _para_in_cell(r2.cells[0],
        "Apakah substansi laporan sudah pernah atau sedang diajukan gugatan ke pengadilan?", size=9)
    _para_in_cell(r2.cells[1], _checkbox(digugat),     size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(r2.cells[2], _checkbox(not digugat), size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    peng  = _safe(validasi.get("nama_pengadilan"))
    noreg = _safe(validasi.get("nomor_register_perkara"))
    _para_in_cell(r2.cells[3],
        f"Pengadilan              : {peng}\nNomor Register Perkara : {noreg}", size=8)

    doc.add_paragraph()
    _build_kronologi_table(doc, kronologi)
    _build_page_footer(doc, 2)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER — TABEL KRONOLOGI
# ═══════════════════════════════════════════════════════════════════════════════

def _build_kronologi_table(doc, kronologi):
    _add_para(doc, "URAIAN PERISTIWA (KRONOLOGI)", bold=True, size=10,
               align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=2)

    rows_data = []
    if isinstance(kronologi, list):
        for item in kronologi:
            if isinstance(item, dict):
                rows_data.append((
                    _safe(item.get("tanggal")),
                    _safe(item.get("peristiwa")),
                    _safe(item.get("catatan_bukti")),
                ))
    elif isinstance(kronologi, dict):
        rows_data.append((
            _safe(kronologi.get("waktu_peristiwa")),
            _safe(kronologi.get("uraian_peristiwa_lengkap")),
            "",
        ))

    while len(rows_data) < 6:
        rows_data.append(("", "", ""))

    tbl = doc.add_table(rows=1 + len(rows_data), cols=3)
    tbl.style = "Table Grid"
    _set_col_widths(tbl, [3.0, 10.0, 3.5])

    hdr = tbl.rows[0]
    for c in hdr.cells:
        _set_cell_bg(c, "D9D9D9")
        _set_cell_borders(c, True, True, True, True)
        _cell_margins(c, 60, 30, 100, 100)
    _para_in_cell(hdr.cells[0], "Tanggal\n(hh/bb/tttt)", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr.cells[1], "Peristiwa\n(5W 1 H)", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(hdr.cells[2], "Catatan/Bukti", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, (tgl, peristiwa, catatan) in enumerate(rows_data):
        row = tbl.rows[i + 1]
        for c in row.cells:
            _set_cell_borders(c, True, True, True, True)
            _cell_margins(c, 40, 40, 100, 100)
        trPr = row._tr.get_or_add_trPr()
        trHeight = OxmlElement("w:trHeight")
        trHeight.set(qn("w:val"), "800")
        trHeight.set(qn("w:hRule"), "atLeast")
        trPr.append(trHeight)
        _para_in_cell(row.cells[0], tgl,       size=9)
        _para_in_cell(row.cells[1], peristiwa,  size=9)
        _para_in_cell(row.cells[2], catatan,    size=9)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER — HALAMAN 4
# ═══════════════════════════════════════════════════════════════════════════════

def _build_hal4(doc, data: dict):
    akhir = data.get("HAL_4_AKHIR", {}) or {}
    ttd   = data.get("HAL_4_TTD",   {}) or {}

    doc.add_paragraph()
    _add_para(doc, "HARAPAN PELAPOR:", bold=True, size=10, space_before=2, space_after=2)

    harapan = _safe(akhir.get("harapan_pelapor"), "………………………………………………………………………")
    tbl1 = doc.add_table(rows=1, cols=1)
    tbl1.style = "Table Grid"
    _set_col_widths(tbl1, [16.5])
    c = tbl1.rows[0].cells[0]
    _set_cell_borders(c, True, True, True, True)
    _cell_margins(c, 80, 80, 120, 120)
    trPr = tbl1.rows[0]._tr.get_or_add_trPr()
    trH = OxmlElement("w:trHeight")
    trH.set(qn("w:val"), "1200")
    trH.set(qn("w:hRule"), "atLeast")
    trPr.append(trH)
    _para_in_cell(c, harapan, size=9)

    doc.add_paragraph()

    # Dokumen pendukung & identitas dirahasiakan
    tbl2 = doc.add_table(rows=2, cols=4)
    tbl2.style = "Table Grid"
    _set_col_widths(tbl2, [7.0, 2.5, 2.5, 4.5])
    dok = _safe(akhir.get("dokumen_pendukung"), "").lower()
    ada_dok = "ada" in dok and "tidak" not in dok
    rahasia = _safe(akhir.get("permintaan_identitas_dirahasiakan"), "").lower() == "ya"
    for row in tbl2.rows:
        for c in row.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 40, 40, 80, 80)

    r0 = tbl2.rows[0]
    _para_in_cell(r0.cells[0],
        "DOKUMEN PENDUKUNG:\nIdentitas Pelapor (KTP / KITAP / KITAS / Lainnya)", bold=True, size=9)
    _para_in_cell(r0.cells[1], _checkbox(ada_dok) + " Ada",           size=10)
    _para_in_cell(r0.cells[2], _checkbox(not ada_dok) + " Tidak Ada", size=10)
    _para_in_cell(r0.cells[3], "", size=9)

    r1 = tbl2.rows[1]
    _para_in_cell(r1.cells[0], "Permintaan Identitas Dirahasiakan:", bold=True, size=9)
    _para_in_cell(r1.cells[1], _checkbox(rahasia) + " Ya",       size=10)
    _para_in_cell(r1.cells[2], _checkbox(not rahasia) + " Tidak", size=10)
    _para_in_cell(r1.cells[3], "", size=9)

    doc.add_paragraph()
    _add_para(doc, "CATATAN:", bold=True, size=10, space_before=2, space_after=2)
    catatan = _safe(akhir.get("catatan_tambahan"), "")
    tbl3 = doc.add_table(rows=1, cols=1)
    tbl3.style = "Table Grid"
    _set_col_widths(tbl3, [16.5])
    c3 = tbl3.rows[0].cells[0]
    _set_cell_borders(c3, True, True, True, True)
    _cell_margins(c3, 80, 80, 120, 120)
    trPr3 = tbl3.rows[0]._tr.get_or_add_trPr()
    trH3 = OxmlElement("w:trHeight")
    trH3.set(qn("w:val"), "1000")
    trH3.set(qn("w:hRule"), "atLeast")
    trPr3.append(trH3)
    _para_in_cell(c3, catatan, size=9)

    doc.add_paragraph()

    # Tanda tangan
    kota_tgl   = _safe(ttd.get("kota_tanggal_ttd"), "Kendari, ………………………… 20 …..")
    nm_pelapor = _safe(ttd.get("nama_pelapor_ttd"))
    nm_petugas = _safe(ttd.get("nama_penerima_laporan_petugas"))

    tbl4 = doc.add_table(rows=4, cols=3)
    tbl4.style = "Table Grid"
    _set_col_widths(tbl4, [5.5, 6.0, 5.0])
    for row in tbl4.rows:
        for c in row.cells:
            _set_cell_borders(c, False, False, False, False)
            _cell_margins(c, 40, 40, 80, 80)

    _para_in_cell(tbl4.rows[0].cells[0], "", size=9)
    _para_in_cell(tbl4.rows[0].cells[1], "", size=9)
    _para_in_cell(tbl4.rows[0].cells[2], kota_tgl, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(tbl4.rows[1].cells[0], "Pelapor", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(tbl4.rows[1].cells[1], "", size=9)
    _para_in_cell(tbl4.rows[1].cells[2], "Penerima Laporan", bold=True, size=9,
                   align=WD_ALIGN_PARAGRAPH.CENTER)
    row2 = tbl4.rows[2]
    trPr2 = row2._tr.get_or_add_trPr()
    trH2 = OxmlElement("w:trHeight")
    trH2.set(qn("w:val"), "1200")
    trH2.set(qn("w:hRule"), "atLeast")
    trPr2.append(trH2)
    for c in row2.cells:
        _para_in_cell(c, "", size=9)
    _para_in_cell(tbl4.rows[3].cells[0], nm_pelapor, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
    _para_in_cell(tbl4.rows[3].cells[1], "", size=9)
    _para_in_cell(tbl4.rows[3].cells[2], nm_petugas, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)

    doc.add_paragraph()
    _add_para(doc, "Catatan: Dokumen yang telah diserahkan menjadi milik Ombudsman RI",
               size=8, space_before=2, space_after=2)

    _build_page_footer(doc, 4)


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER — FOOTER HALAMAN
# ═══════════════════════════════════════════════════════════════════════════════

def _build_page_footer(doc, page_num: int):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    top = OxmlElement("w:top")
    top.set(qn("w:val"), "single")
    top.set(qn("w:sz"), "6")
    top.set(qn("w:color"), "000000")
    pBdr.append(top)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(
        f"{page_num} | Unit Penerimaan dan Verifikasi Laporan (PVL)\n"
        "OMBUDSMAN RI PERWAKILAN PROVINSI SULTRA"
    )
    run.font.size = Pt(8)
    run.bold = True
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT


# ═══════════════════════════════════════════════════════════════════════════════
# DOCX BUILDER — FUNGSI UTAMA
# ═══════════════════════════════════════════════════════════════════════════════

def create_ombudsman_docx(forms: list, output_path: str):
    """Buat dokumen Word dengan satu formulir penuh (4 halaman) per item di `forms`."""
    # ── Gunakan resource_path() agar logo terbaca saat di-bundle PyInstaller ──
    logo_left  = resource_path("logo-kiri1.jpg")
    logo_right = resource_path("logo-kanan.png")

    doc = Document()

    # Setup halaman A4
    section = doc.sections[0]
    section.page_width    = Cm(21.0)
    section.page_height   = Cm(29.7)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.0)
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(10)

    for form_idx, form in enumerate(forms):
        if form_idx > 0:
            _page_break(doc)

        # Hal 1
        _build_header(doc, logo_left, logo_right)
        _build_hal1(doc, form)
        _page_break(doc)

        # Hal 2
        _build_hal2(doc, form)
        _page_break(doc)

        # Hal 3 – lanjutan kronologi (kosong)
        _add_para(doc, "URAIAN PERISTIWA (KRONOLOGI) – Lanjutan", bold=True,
                   size=10, align=WD_ALIGN_PARAGRAPH.CENTER, space_before=2, space_after=4)
        empty_kron = [("", "", "") for _ in range(8)]
        tbl = doc.add_table(rows=1 + len(empty_kron), cols=3)
        tbl.style = "Table Grid"
        _set_col_widths(tbl, [3.0, 10.0, 3.5])
        hdr = tbl.rows[0]
        for c in hdr.cells:
            _set_cell_bg(c, "D9D9D9")
            _set_cell_borders(c, True, True, True, True)
            _cell_margins(c, 60, 30, 100, 100)
        _para_in_cell(hdr.cells[0], "Tanggal\n(hh/bb/tttt)", bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _para_in_cell(hdr.cells[1], "Peristiwa\n(5W 1 H)", bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _para_in_cell(hdr.cells[2], "Catatan/Bukti", bold=True, size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        for i in range(len(empty_kron)):
            row = tbl.rows[i + 1]
            for c in row.cells:
                _set_cell_borders(c, True, True, True, True)
                _cell_margins(c, 40, 40, 100, 100)
            trPr = row._tr.get_or_add_trPr()
            trH = OxmlElement("w:trHeight")
            trH.set(qn("w:val"), "800")
            trH.set(qn("w:hRule"), "atLeast")
            trPr.append(trH)
            for c in row.cells:
                _para_in_cell(c, "", size=9)

        _build_page_footer(doc, 3)
        _page_break(doc)

        # Hal 4
        _build_hal4(doc, form)

    doc.save(output_path)

    # Bersihkan pengaturan zoom yang tidak valid dari python-docx
    tmp = output_path + ".tmp"
    _sh.copy(output_path, tmp)
    with _zf.ZipFile(tmp, "r") as zin, _zf.ZipFile(output_path, "w", _zf.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/settings.xml":
                tree = etree.fromstring(data)
                ns = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
                for z in tree.findall(f"{{{ns}}}zoom"):
                    tree.remove(z)
                data = etree.tostring(tree, xml_declaration=True, encoding="UTF-8", standalone=True)
            zout.writestr(item, data)
    os.remove(tmp)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
# MULTI-FORM SPLITTER
# ═══════════════════════════════════════════════════════════════════════════════

def split_pages_into_form_groups(all_pages: list) -> list:
    """
    Kelompokkan halaman menjadi grup per formulir.

    Aturan:
    - 1 formulir = PAGES_PER_FORM (4) halaman
    - Jika total halaman bukan kelipatan 4, sisa halaman tetap
      diproses sebagai satu formulir (formulir tidak lengkap)
    - Contoh: 9 halaman → grup [4, 4, 1]

    Returns:
        list of list[PIL.Image] — setiap sub-list = 4 halaman satu formulir
    """
    groups = []
    total  = len(all_pages)
    for start in range(0, total, PAGES_PER_FORM):
        chunk = all_pages[start : start + PAGES_PER_FORM]
        groups.append(chunk)
    return groups


def detect_form_count(total_pages: int) -> tuple:
    """
    Hitung jumlah formulir dan beri peringatan jika tidak pas.

    Returns:
        (jumlah_formulir, pesan_peringatan | "")
    """
    full_forms  = total_pages // PAGES_PER_FORM
    sisa        = total_pages %  PAGES_PER_FORM
    warning_msg = ""

    if total_pages == 0:
        return 0, "Tidak ada halaman yang dimuat!"

    if sisa != 0:
        warning_msg = (
            f"⚠  Total halaman ({total_pages}) bukan kelipatan {PAGES_PER_FORM}.\n"
            f"   Ditemukan {full_forms} formulir lengkap + {sisa} halaman sisa.\n"
            f"   Halaman sisa akan tetap diproses sebagai formulir tidak lengkap."
        )
        full_forms += 1   # proses juga sisa halaman

    return full_forms, warning_msg


# ═══════════════════════════════════════════════════════════════════════════════
# GUI — TKINTER APP
# ═══════════════════════════════════════════════════════════════════════════════

def _load_logo(path: str, size: tuple):
    try:
        img = Image.open(path).convert("RGBA")
        img = img.resize(size, Image.LANCZOS)
        return ImageTk.PhotoImage(img)
    except Exception:
        return None


class App:
    def __init__(self, root):
        self.root = root
        self.root.title("Ombudsman → Word Generator · v2.1")
        self.root.geometry("540x720")
        self.root.minsize(480, 640)
        self.root.configure(bg=C["bg"])
        self.root.resizable(True, True)

        self.api_key    = None
        self.model_name = "gemini-2.5-flash"
        self.client     = None
        self.file_paths = []
        self._pulse_job = None
        self._dot_count = 0
        self._logo_left_img  = None
        self._logo_right_img = None

        # State multi-form
        self._all_pages      = []   # semua PIL.Image yang sudah dimuat
        self._form_groups    = []   # [[page, page, page, page], ...]
        self._extracted_forms = []  # hasil JSON per formulir

        self._load_config()
        self._build_ui()

    # ── Konfigurasi ──────────────────────────────────────────
    def _load_config(self):
        if os.path.exists(CONFIG_FILE):
            try:
                with open(CONFIG_FILE) as f:
                    cfg = json.load(f)
                    self.api_key    = cfg.get("api_key")
                    self.model_name = cfg.get("model_name", "gemini-2.5-flash")
            except Exception:
                pass
        if not self.api_key:
            self._show_setup()
        else:
            self._init_client()

    def _init_client(self):
        try:
            self.client = genai.Client(api_key=self.api_key)
        except Exception as e:
            messagebox.showerror("Error", f"Gagal menghubungkan ke Gemini API:\n{e}")

    def _show_setup(self):
        win = tk.Toplevel(self.root)
        win.title("Konfigurasi API")
        win.geometry("440x320")
        win.configure(bg=C["bg"])
        win.grab_set()
        win.resizable(False, False)

        tk.Frame(win, bg=C["accent"], height=3).pack(fill="x")
        frm = tk.Frame(win, bg=C["bg"], padx=28, pady=22)
        frm.pack(fill="both", expand=True)

        tk.Label(frm, text="⚙  Konfigurasi Gemini API", font=("Segoe UI", 12, "bold"),
                 bg=C["bg"], fg=C["text_primary"]).pack(anchor="w")
        tk.Label(frm, text="Masukkan API Key dan pilih model", font=FONT_SMALL,
                 bg=C["bg"], fg=C["text_secondary"]).pack(anchor="w", pady=(2, 16))

        tk.Label(frm, text="API KEY", font=("Segoe UI", 8, "bold"),
                 bg=C["bg"], fg=C["text_muted"]).pack(anchor="w")
        entry = tk.Entry(frm, font=FONT_BODY, bg=C["surface2"], fg=C["text_primary"],
                         relief="flat", highlightthickness=1,
                         highlightbackground=C["border"], highlightcolor=C["accent"],
                         insertbackground=C["accent"])
        entry.pack(fill="x", ipady=7, pady=(4, 14))

        tk.Label(frm, text="MODEL", font=("Segoe UI", 8, "bold"),
                 bg=C["bg"], fg=C["text_muted"]).pack(anchor="w")
        models = ["gemini-2.5-flash", "gemini-2.0-flash", "gemini-1.5-pro"]
        mv = tk.StringVar(value=models[0])
        ttk.Combobox(frm, textvariable=mv, values=models, state="readonly",
                     font=FONT_BODY).pack(fill="x", ipady=4, pady=(4, 18))

        def save():
            k = entry.get().strip()
            if not k:
                messagebox.showwarning("Peringatan", "API Key wajib diisi!", parent=win)
                return
            cfg = {"api_key": k, "model_name": mv.get()}
            with open(CONFIG_FILE, "w") as f:
                json.dump(cfg, f)
            self.api_key    = k
            self.model_name = mv.get()
            self._init_client()
            win.destroy()

        tk.Button(frm, text="Simpan & Mulai →", font=FONT_BTN,
                  bg=C["accent"], fg=C["white"], relief="flat",
                  activebackground=C["accent_dark"], activeforeground=C["white"],
                  cursor="hand2", bd=0, command=save).pack(fill="x", ipady=10)
        win.protocol("WM_DELETE_WINDOW",
                     lambda: (self.root.destroy() if not self.api_key else win.destroy()))

    # ── UI ───────────────────────────────────────────────────
    def _build_ui(self):
        s = ttk.Style()
        s.theme_use("clam")
        s.configure("P.Horizontal.TProgressbar", troughcolor=C["surface2"],
                    background=C["accent"], thickness=6)

        wrap = tk.Frame(self.root, bg=C["bg"], padx=22, pady=18)
        wrap.pack(fill="both", expand=True)
        wrap.columnconfigure(0, weight=1)
        r = 0

        # Header dengan logo — gunakan resource_path() untuk GUI juga
        hdr = tk.Frame(wrap, bg=C["bg"])
        hdr.grid(row=r, column=0, sticky="ew", pady=(0, 14)); r += 1
        hdr.columnconfigure(1, weight=1)

        self._logo_left_img = _load_logo(LOGO_LEFT_PATH, LOGO_SIZE)
        if self._logo_left_img:
            tk.Label(hdr, image=self._logo_left_img, bg=C["bg"]).grid(
                row=0, column=0, rowspan=2, padx=(0, 12), sticky="w")

        txt_frame = tk.Frame(hdr, bg=C["bg"])
        txt_frame.grid(row=0, column=1, rowspan=2, sticky="nsew")
        tk.Label(txt_frame, text="📄  Ombudsman → Word Generator",
                 font=("Segoe UI", 13, "bold"), bg=C["bg"],
                 fg=C["text_primary"]).pack(anchor="w")
        tk.Label(txt_frame,
                 text="Scan formulir PDF → ekstrak data OCR → ekspor ke Word (.docx)",
                 font=FONT_SMALL, bg=C["bg"], fg=C["text_secondary"]).pack(anchor="w", pady=(2, 0))

        self._logo_right_img = _load_logo(LOGO_RIGHT_PATH, LOGO_SIZE)
        if self._logo_right_img:
            tk.Label(hdr, image=self._logo_right_img, bg=C["bg"]).grid(
                row=0, column=2, rowspan=2, padx=(12, 0), sticky="e")

        tk.Frame(wrap, bg=C["border"], height=1).grid(row=r, column=0, sticky="ew", pady=(0, 14)); r += 1

        # Card file
        card = tk.Frame(wrap, bg=C["surface"], highlightthickness=1,
                        highlightbackground=C["border"])
        card.grid(row=r, column=0, sticky="ew", pady=(0, 10)); r += 1
        card.columnconfigure(0, weight=1)
        tk.Frame(card, bg=C["accent"], height=2).grid(row=0, column=0, sticky="ew")
        inner = tk.Frame(card, bg=C["surface"], padx=16, pady=12)
        inner.grid(row=1, column=0, sticky="ew")
        inner.columnconfigure(0, weight=1)
        tk.Label(inner, text="BERKAS INPUT", font=("Segoe UI", 8, "bold"),
                 bg=C["surface"], fg=C["text_muted"]).grid(row=0, column=0, sticky="w")

        self.lbl_files = tk.Label(inner,
            text="Belum ada file dipilih\nKlik tombol di bawah untuk memilih PDF / gambar",
            font=FONT_SMALL, bg=C["surface2"], fg=C["text_muted"],
            justify="center", pady=20, anchor="center",
            highlightthickness=1, highlightbackground=C["border"])
        self.lbl_files.grid(row=1, column=0, sticky="ew", pady=(6, 10))

        tk.Button(inner, text="＋  Pilih File  (JPG / PNG / PDF)",
                  font=FONT_BTN, bg=C["surface2"], fg=C["accent"],
                  relief="flat", cursor="hand2", bd=0,
                  highlightthickness=1, highlightbackground=C["border"],
                  activebackground=C["border"], activeforeground=C["accent_dark"],
                  command=self._select_files).grid(row=2, column=0, sticky="ew", ipady=8)

        # ── Info card: formulir terdeteksi ──────────────────
        info_card = tk.Frame(wrap, bg=C["surface"], highlightthickness=1,
                              highlightbackground=C["border"])
        info_card.grid(row=r, column=0, sticky="ew", pady=(0, 10)); r += 1
        info_card.columnconfigure(0, weight=1)
        tk.Frame(info_card, bg=C["warning"], height=2).grid(row=0, column=0, sticky="ew")
        info_inner = tk.Frame(info_card, bg=C["surface"], padx=16, pady=10)
        info_inner.grid(row=1, column=0, sticky="ew")
        info_inner.columnconfigure(1, weight=1)

        tk.Label(info_inner, text="DETEKSI FORMULIR", font=("Segoe UI", 8, "bold"),
                 bg=C["surface"], fg=C["text_muted"]).grid(
                 row=0, column=0, columnspan=3, sticky="w", pady=(0, 6))

        # Baris info
        self._lbl_total_pages = tk.Label(info_inner, text="Total halaman : –",
                                          font=FONT_SMALL, bg=C["surface"],
                                          fg=C["text_secondary"])
        self._lbl_total_pages.grid(row=1, column=0, sticky="w", padx=(0, 20))

        self._lbl_form_count = tk.Label(info_inner, text="Formulir      : –",
                                         font=("Segoe UI", 9, "bold"), bg=C["surface"],
                                         fg=C["accent"])
        self._lbl_form_count.grid(row=1, column=1, sticky="w")

        self._lbl_pages_per = tk.Label(info_inner,
                                        text=f"(1 formulir = {PAGES_PER_FORM} halaman)",
                                        font=FONT_SMALL, bg=C["surface"],
                                        fg=C["text_muted"])
        self._lbl_pages_per.grid(row=1, column=2, sticky="e")

        self._lbl_warn = tk.Label(info_inner, text="", font=FONT_SMALL,
                                   bg=C["surface"], fg=C["warning"],
                                   justify="left", wraplength=440)
        self._lbl_warn.grid(row=2, column=0, columnspan=3, sticky="w", pady=(4, 0))

        # ── Override jumlah formulir ────────────────────────
        override_frame = tk.Frame(info_inner, bg=C["surface"])
        override_frame.grid(row=3, column=0, columnspan=3, sticky="w", pady=(6, 0))
        tk.Label(override_frame, text="Override jumlah formulir (opsional):",
                 font=FONT_SMALL, bg=C["surface"], fg=C["text_secondary"]).pack(side="left")
        self._override_var = tk.StringVar(value="")
        tk.Entry(override_frame, textvariable=self._override_var,
                 font=FONT_BODY, width=5, bg=C["surface2"], fg=C["text_primary"],
                 relief="flat", highlightthickness=1,
                 highlightbackground=C["border"]).pack(side="left", padx=(8, 0), ipady=3)
        tk.Label(override_frame, text="(kosongkan = otomatis)",
                 font=FONT_SMALL, bg=C["surface"], fg=C["text_muted"]).pack(side="left", padx=(6, 0))

        # Badge model
        mb = tk.Frame(wrap, bg=C["bg"])
        mb.grid(row=r, column=0, sticky="ew", pady=(0, 10)); r += 1
        mb.columnconfigure(1, weight=1)
        dot = tk.Canvas(mb, width=8, height=8, bg=C["bg"], highlightthickness=0)
        dot.grid(row=0, column=0, padx=(0, 6), pady=3)
        dot.create_oval(1, 1, 7, 7, fill=C["success"], outline="")
        self.lbl_model = tk.Label(mb, text=f"Model: {self.model_name}",
                                   font=FONT_SMALL, bg=C["bg"], fg=C["text_secondary"])
        self.lbl_model.grid(row=0, column=1, sticky="w")
        lnk = tk.Label(mb, text="⚙ Ganti", font=FONT_SMALL, bg=C["bg"],
                        fg=C["accent"], cursor="hand2")
        lnk.grid(row=0, column=2, sticky="e")
        lnk.bind("<Button-1>", lambda e: self._show_setup())

        # Tombol proses
        self.btn = tk.Button(wrap, text="⚡  Proses & Buat Word (.docx)",
                              font=("Segoe UI", 11, "bold"),
                              bg=C["accent"], fg=C["white"],
                              activebackground=C["accent_dark"], activeforeground=C["white"],
                              relief="flat", cursor="hand2", bd=0,
                              command=self._start)
        self.btn.grid(row=r, column=0, sticky="ew", ipady=13, pady=(0, 10)); r += 1

        # Status
        sfrm = tk.Frame(wrap, bg=C["surface"], highlightthickness=1,
                         highlightbackground=C["border"])
        sfrm.grid(row=r, column=0, sticky="ew", pady=(0, 8)); r += 1
        sfrm.columnconfigure(1, weight=1)
        si = tk.Frame(sfrm, bg=C["surface"], padx=12, pady=10)
        si.grid(row=0, column=0, sticky="ew")
        si.columnconfigure(1, weight=1)
        self._sdot = tk.Canvas(si, width=8, height=8, bg=C["surface"], highlightthickness=0)
        self._sdot.grid(row=0, column=0, padx=(0, 8))
        self._sdot_oval = self._sdot.create_oval(1, 1, 7, 7, fill=C["text_muted"], outline="")
        self._svar = tk.StringVar(value="Sistem siap · Pilih file untuk memulai")
        tk.Label(si, textvariable=self._svar, font=FONT_SMALL,
                 bg=C["surface"], fg=C["text_secondary"], anchor="w").grid(row=0, column=1, sticky="w")

        # Progress bar
        self._pvar = tk.DoubleVar(value=0)
        self._pbar = ttk.Progressbar(wrap, variable=self._pvar, maximum=100,
                                      mode="determinate", style="P.Horizontal.TProgressbar")
        self._pbar.grid(row=r, column=0, sticky="ew", pady=(0, 18)); r += 1
        self._pbar.config(length=1)

        # Footer
        tk.Frame(wrap, bg=C["border"], height=1).grid(row=r, column=0, sticky="ew", pady=(4, 8)); r += 1
        tk.Label(wrap,
                 text="© 2026  Ombudsman RI Perwakilan Sulawesi Tenggara || Universitas Halu Oleo\n"
                      "Abrar Wujedaan dan Abdul Mu'iz Azizul Raeba",
                 font=("Segoe UI", 7), bg=C["bg"], fg=C["text_muted"],
                 justify="center").grid(row=r, column=0)

    # ── Helpers ──────────────────────────────────────────────
    def _set_status(self, msg, color=None):
        self._svar.set(msg)
        self._sdot.itemconfig(self._sdot_oval, fill=color or C["text_muted"])

    def _set_progress(self, pct: float):
        """Set progress bar 0–100, tampilkan jika > 0."""
        if pct <= 0:
            self._pvar.set(0)
            self._pbar.config(length=1)
        else:
            self._pbar.config(length=400)
            self._pvar.set(min(pct, 100))

    def _pulse(self):
        cols = [C["accent"], "#3B82F6", C["accent_dark"]]
        self._dot_count = (self._dot_count + 1) % len(cols)
        self._sdot.itemconfig(self._sdot_oval, fill=cols[self._dot_count])
        self._pulse_job = self.root.after(400, self._pulse)

    def _stop_pulse(self):
        if self._pulse_job:
            self.root.after_cancel(self._pulse_job)
            self._pulse_job = None

    # ── Pilih file ───────────────────────────────────────────
    def _select_files(self):
        files = filedialog.askopenfilenames(
            title="Pilih File Formulir Ombudsman",
            filetypes=[("Gambar & PDF", "*.jpg *.jpeg *.png *.pdf"), ("Semua", "*.*")])
        if not files:
            return
        self.file_paths = list(files)
        n = len(self.file_paths)

        # Preview nama file
        if n <= 5:
            txt = "\n".join(f"  ✓  {os.path.basename(p)}" for p in self.file_paths)
        else:
            txt = "\n".join(f"  ✓  {os.path.basename(p)}" for p in self.file_paths[:4])
            txt += f"\n  … dan {n-4} file lainnya"
        self.lbl_files.config(text=txt, fg=C["text_primary"], justify="left",
                               anchor="w", pady=12)

        # Hitung halaman dan formulir (jalankan di thread agar tidak freeze)
        self._lbl_total_pages.config(text="Total halaman : menghitung…")
        self._lbl_form_count.config(text="Formulir      : …", fg=C["text_muted"])
        self._lbl_warn.config(text="")
        self._set_status("Memuat & menghitung halaman…", C["accent"])
        threading.Thread(target=self._precount_pages, daemon=True).start()

    def _precount_pages(self):
        """Hitung total halaman di background tanpa konversi penuh."""
        try:
            total = 0
            for p in self.file_paths:
                if p.lower().endswith(".pdf"):
                    # convert_from_path hanya untuk menghitung — gunakan dpi rendah
                    pages = convert_from_path(p, dpi=72)
                    total += len(pages)
                else:
                    total += 1   # gambar = 1 halaman

            jumlah, warn = detect_form_count(total)

            def _update():
                self._lbl_total_pages.config(
                    text=f"Total halaman : {total}")
                self._lbl_form_count.config(
                    text=f"Formulir      : {jumlah}",
                    fg=C["accent"] if not warn else C["warning"])
                self._lbl_warn.config(text=warn)
                self._set_status(
                    f"{total} halaman · {jumlah} formulir terdeteksi", C["success"])

            self.root.after(0, _update)

        except Exception as e:
            err_msg = str(e)
            self.root.after(0, lambda m=err_msg: self._set_status(
                f"Gagal hitung halaman: {m[:60]}", C["danger"]))

    # ── Proses ───────────────────────────────────────────────
    def _start(self):
        if not self.file_paths:
            self._set_status("⚠ Pilih file terlebih dahulu!", C["warning"])
            return
        if not self.client:
            messagebox.showerror("Error", "Client AI tidak aktif. Cek API Key.")
            return
        self.btn.config(state="disabled", text="Memproses…", bg=C["surface2"])
        self._set_progress(1)  # tampilkan progress bar
        self._pulse()
        self._extracted_forms = []
        threading.Thread(target=self._run, daemon=True).start()

    def _run(self):
        try:
            # ── 1. Muat semua gambar ────────────────────────
            self.root.after(0, lambda: self._set_status("Memuat gambar…", C["accent"]))
            self._all_pages = self._load_images()
            total = len(self._all_pages)

            # ── 2. Tentukan jumlah formulir ─────────────────
            override_str = self._override_var.get().strip()
            if override_str.isdigit() and int(override_str) > 0:
                n_forms = int(override_str)
                # Bagi rata halaman ke n_forms grup
                pages_each = total // n_forms
                rem = total % n_forms
                groups = []
                idx = 0
                for i in range(n_forms):
                    extra = 1 if i < rem else 0
                    end = idx + pages_each + extra
                    groups.append(self._all_pages[idx:end])
                    idx = end
                self._form_groups = groups
                warn_msg = ""
            else:
                n_forms, warn_msg = detect_form_count(total)
                self._form_groups = split_pages_into_form_groups(self._all_pages)

            if warn_msg:
                self.root.after(0, lambda w=warn_msg: self._lbl_warn.config(text=w))

            self.root.after(0, lambda n=n_forms: self._lbl_form_count.config(
                text=f"Formulir      : {n}", fg=C["accent"]))
            self.root.after(0, lambda t=total: self._lbl_total_pages.config(
                text=f"Total halaman : {t}"))

            # ── 3. Ekstrak tiap grup formulir satu per satu ─
            self._extracted_forms = []
            for i, group in enumerate(self._form_groups):
                form_num = i + 1
                msg = f"Mengekstrak formulir {form_num}/{n_forms} ({len(group)} hal)…"
                self.root.after(0, lambda m=msg: self._set_status(m, C["accent"]))
                pct = (i / n_forms) * 90 + 5   # 5% – 95%
                self.root.after(0, lambda p=pct: self._set_progress(p))

                form_data = self._extract_one_form(group, form_num, n_forms)
                self._extracted_forms.append(form_data)

            # ── 4. Simpan ───────────────────────────────────
            self.root.after(0, lambda: self._set_status("Selesai. Menyimpan…", C["success"]))
            self.root.after(0, lambda: self._set_progress(97))

            extracted = self._extracted_forms
            n = n_forms
            self.root.after(0, lambda f=extracted, n=n: self._ask_save(f, n))

        except Exception as e:
            import traceback
            tb  = traceback.format_exc()
            err_msg = str(e)
            err_tb  = tb
            self.root.after(0, lambda m=err_msg, t=err_tb: messagebox.showerror(
                "Error", f"{m}\n\n{t}"))
            self.root.after(0, lambda m=err_msg: self._set_status(
                f"Gagal: {m[:70]}", C["danger"]))
        finally:
            self.root.after(0, self._reset_ui)

    def _extract_one_form(self, pages: list, form_num: int, total_forms: int) -> dict:
        """
        Kirim satu grup halaman (satu formulir) ke Gemini dan kembalikan dict data.
        Jika AI menghasilkan format lama (berisi data_formulir), ambil elemen pertama.
        """
        prompt_konteks = (
            f"Ini adalah formulir ke-{form_num} dari total {total_forms} formulir. "
            f"Formulir ini terdiri dari {len(pages)} halaman."
        )

        response = self.client.models.generate_content(
            model=self.model_name,
            contents=[EXTRACTION_PROMPT, prompt_konteks, *pages],
            config={"temperature": 0, "top_p": 1, "top_k": 1}
        )

        result = self._parse_json(response.text)

        # Kompatibilitas: jika AI salah kembalikan format lama (list di data_formulir)
        if "data_formulir" in result and isinstance(result["data_formulir"], list):
            forms_list = result["data_formulir"]
            if forms_list:
                return forms_list[0]
            return {}

        return result

    def _load_images(self):
        images = []
        for p in self.file_paths:
            if p.lower().endswith(".pdf"):
                pages = convert_from_path(p, dpi=250)
                images.extend(pages)
            else:
                img = Image.open(p)
                if img.mode not in ("RGB", "L"):
                    img = img.convert("RGB")
                images.append(img)
        return images

    def _parse_json(self, raw: str) -> dict:
        clean = re.sub(r'```(?:json)?\s*', '', raw).strip().rstrip('`').strip()
        try:
            return json.loads(clean)
        except json.JSONDecodeError:
            pass
        m = re.search(r'\{[\s\S]*\}', clean)
        if m:
            try:
                return json.loads(m.group())
            except Exception:
                pass
        raise ValueError(f"Tidak bisa parse JSON:\n{raw[:400]}")

    def _ask_save(self, forms, jumlah):
        first  = forms[0] if forms else {}
        nama   = first.get("HAL_1_PELAPOR", {}).get("nama_lengkap", "Unknown") or "Unknown"
        default = f"Ombudsman_{jumlah}_Formulir_{nama.replace(' ', '_')}.docx"
        path = filedialog.asksaveasfilename(
            defaultextension=".docx",
            filetypes=[("Word Document", "*.docx")],
            initialfile=default)
        if not path:
            self._set_status("Dibatalkan · File tidak disimpan.", C["warning"])
            return
        try:
            self._set_status("Membuat file Word…", C["accent"])
            create_ombudsman_docx(forms, path)
            self._set_progress(100)
            self._set_status(f"✓ {jumlah} formulir disimpan sebagai Word", C["success"])
            messagebox.showinfo("Berhasil",
                f"✅  {jumlah} formulir berhasil diekspor!\n\n📁  {path}")
        except Exception as e:
            err_msg = str(e)
            messagebox.showerror("Gagal Menyimpan", err_msg)
            self.root.after(0, lambda m=err_msg: self._set_status(
                f"Gagal simpan: {m[:60]}", C["danger"]))

    def _reset_ui(self):
        self._stop_pulse()
        self._set_progress(0)
        self.btn.config(state="normal", text="⚡  Proses & Buat Word (.docx)", bg=C["accent"])


# ═══════════════════════════════════════════════════════════════════════════════
# ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    root = tk.Tk()
    App(root)
    root.mainloop()